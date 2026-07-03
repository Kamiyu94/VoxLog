import os
import re
import sys
import json
import time
import math
import platform
import threading
import subprocess
import shutil
import multiprocessing
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

if platform.system() == "Windows":
    # 把 winget 安裝的 ffmpeg 加進 PATH。原本寫死成特定使用者的路徑，
    # 換一台電腦（不同使用者名稱／ffmpeg 版本）就失效，這裡改成自動尋找。
    import glob
    _ff_pattern = os.path.join(
        os.environ.get("LOCALAPPDATA", ""),
        "Microsoft", "WinGet", "Packages", "Gyan.FFmpeg*", "**", "bin",
    )
    for _d in glob.glob(_ff_pattern, recursive=True):
        if os.path.isfile(os.path.join(_d, "ffmpeg.exe")):
            os.environ["PATH"] += os.pathsep + _d
            break

FONT_UI = "Microsoft JhengHei" if platform.system() == "Windows" else "PingFang TC"

CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")

BG        = "#1E1E1E"
SURFACE   = "#2A2A2A"
BORDER    = "#3A3A3A"
TEXT      = "#E8E8E8"
SUBTEXT   = "#888888"
ACCENT    = "#4A9EFF"
GREEN     = "#27AE60"
GREEN_DIM = "#1A5C38"
RED       = "#E74C3C"
RED_HOVER = "#C0392B"
BLUE      = "#2471A3"
BLUE_DIM  = "#1A4F72"

VIDEO_EXTS = {".mp4", ".mkv", ".avi", ".mov", ".webm"}

LANG_MAP = {
    "中文": "zh",
    "英文": "en",
    "日文": "ja",
    "韓文": "ko",
    "自動偵測": None,
}

AI_SRT_PROMPT = """以下有 {n} 段逐字稿，每段以 ===SEG編號=== 標記開頭。
請將每段文字依語意切成適合字幕的短句，每句 15 字以內，以自然語意斷句。

輸出格式（嚴格遵守）：
- 每段用 ===SEG編號=== 開頭（與輸入編號對應）
- 短句各占一行
- 不要加任何其他內容、編號或符號

{segments}"""

CORRECT_PROMPT = """以下是一份語音辨識產生的逐字稿，可能含有空耳錯誤、同音異字或辨識錯誤。{context}
請修正明顯的辨識錯誤，規則如下：
- 保留所有說話人標記（如 SPEAKER_00：）和時間戳格式（如 [00:01 --> 00:21]）
- 不要增加、刪除或合併段落
- 不要改變說話內容的意思
- 以繁體中文輸出
- 直接輸出修正後的逐字稿，不要在前後加任何說明文字或前言

逐字稿內容：
{transcript}"""

NOTES_PROMPT = """你是專業的會議記錄整理者。以下是一份逐字稿（可能是會議、訪談、教學或影片內容）。{context}

請產出一份「結構清楚、敘述完整、可直接交付」的繁體中文摘要，品質對標專業會議紀錄工具。要求：
- 不要只丟零碎條列；每個主題段落要有完整敘述（2–5 句），把來龍去脈、原因、結論講清楚。
- 若逐字稿有 SPEAKER 或人名，適當標記是誰說的、誰負責什麼。
- 只根據逐字稿內容，不要杜撰；資訊不足處寫「（逐字稿未明確說明）」。
- 務必保留逐字稿中出現的「具體事實」，不可在摘要時被抽象化或省略，包括：人名與其角色/負責事項、日期與時程節點（截止日、上線日等）、數字與金額、系統與技術名稱（如 CDP、API、Notion）、案號或專案編號、檔案名稱或位置、誰需要提供或確認什麼。寧可摘要長一點，也不要漏掉這些落地細節。
- 轉述原則、指示或結論時，必須保留其正確的語意與方向，不可簡化到語意相反或模糊（例：「要避免 X」不可寫成「要落實 X」）。

嚴格使用以下 Markdown 格式輸出（標題、清單、表格都要照格式）。
注意：不要自己加最上層的 # 大標題（標題會由系統用檔名填入）。

## 基本資訊
- 會議時間：{meeting_time}
- 參與者：（依逐字稿列出說話者或人名；無法判斷就寫「未明確」）

## 會議總覽
（用一段 150–250 字，把整場的背景、核心重點、結論與目標濃縮成流暢的一段話。）

## 會議重點
（3–5 個最關鍵的結論或重點，每個一行條列，簡短有力；聚焦最重要的，不要與會議總覽逐句重複。）
- （重點一）
- （重點二）

## 會議要素
（把這場會議的關鍵要素逐項列出，每項盡量標出處說話者；某類別無資訊則寫「（逐字稿未明確說明）」。）
- 關鍵人物與角色：（誰、負責什麼）
- 時程與期限：（日期、里程碑、deadline）
- 數字與金額：（如有）
- 系統／技術／工具：（提到的系統、API、平台等）
- 案號／專案編號／檔案位置：（如有）

## 討論內容
（依主題切分，每個主題一個 ### 小標；小標後接一段完整敘述，必要時再用 - 補關鍵細節。要涵蓋討論的主要段落。）

### （主題一的標題）
（敘述段落）
- （關鍵細節，可省略）

### （主題二的標題）
（敘述段落）

## 決議事項
（本場確立的決定、共識或工作原則；若無則寫「無」。）
- （條列）

## 待辦事項
（用表格列出每一項待辦；事項中若涉及特定人、檔案、系統，要一併寫出，不要只寫抽象動作；沒有明確期限就寫 TBD。）

| 負責人 | 事項 | 期限 |
|--------|------|------|
| （誰） | （要做什麼，含相關人/檔案/系統） | （日期或 TBD） |

## 標籤

### 人物標籤
（逐字稿中出現或被提到的人名，格式 #人名，同一行以空格分隔；若無則寫「無」）

### 主題標籤
（3–5 個主要主題，格式 #主題，同一行以空格分隔）

### 議題標籤
（5–8 個具體討論的議題或關鍵詞，格式 #議題，同一行以空格分隔）

逐字稿內容：
{transcript}"""

CORRECT_AND_PARTY_PROMPT = """以下是一份語音辨識產生的逐字稿，沒有說話人標記。{context}
請一次完成兩件事：（1）修正明顯的辨識錯誤（空耳、同音異字）；（2）判斷每一段是「哪一方」說的，在每段前面標上方別。規則如下：
- 先判斷這場對話有哪幾方；若上面的背景資訊有指明各方是誰（人名或角色），就用背景給的名稱，否則自行判斷合適的角色名（例如：提供方／諮詢方、主管／部屬、我方／客方）。
- 每一行格式為「〔方別〕」接原本的時間戳與（已校正的）內容，例如：〔王經理〕[00:12 --> 00:18] 我們的產品可以做到…
- 「對」「嗯嗯」「了解」「沒錯」這種沒有實質內容的短附和，歸給「前一位有實質發言的人」，不要自成一方，也不要誤判成另一方。
- 真的完全無法判斷時才標〔不確定〕。
- 保留所有時間戳；不要增加、刪除、合併段落；不要改變說話內容的意思；以繁體中文輸出。
- 直接輸出處理後的逐字稿，前後不要加任何說明或前言。

逐字稿內容：
{transcript}"""


def load_config():
    try:
        with open(CONFIG_PATH, "r") as f:
            return json.load(f)
    except Exception:
        return {}


def save_config(data):
    cfg = load_config()
    cfg.update(data)
    with open(CONFIG_PATH, "w") as f:
        json.dump(cfg, f, indent=2)


def ensure_config():
    """首次啟動若沒有 config.json，自動建立一份空白設定（金鑰留空、保留欄位與預設值），
    使用者完全不必手動建檔，直接在 GUI 填入金鑰按「驗證」即可存檔。"""
    if os.path.exists(CONFIG_PATH):
        return
    template = {}
    example_path = os.path.join(os.path.dirname(CONFIG_PATH), "config.example.json")
    try:
        with open(example_path, "r", encoding="utf-8") as f:
            example = json.load(f)
        for k, v in example.items():
            # placeholder（YOUR_xxx）清空，其餘預設值（如 openai_model）保留
            template[k] = "" if isinstance(v, str) and v.startswith("YOUR_") else v
    except Exception:
        template = {}
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(template, f, indent=2, ensure_ascii=False)
    except Exception:
        pass


def _format_srt_time(seconds):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int(round((seconds - int(seconds)) * 1000))
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _words_to_srt_segments(words, max_chars=20):
    """把逐字時間戳列表依標點/字數切成字幕段落，回傳 [(start, end, text), ...]。"""
    segments = []
    cur_words = []
    cur_text = ""
    for w in words:
        word = w.get("word", "").strip()
        if not word:
            continue
        cur_words.append(w)
        cur_text += word
        ends_sentence = any(c in word for c in "。！？")
        ends_clause = any(c in word for c in "，、") and len(cur_text) >= max_chars
        over_limit = len(cur_text) >= max_chars * 2  # 無標點強制切割
        if (ends_sentence or ends_clause or over_limit) and cur_words:
            segments.append((
                cur_words[0].get("start", 0),
                cur_words[-1].get("end", 0),
                cur_text,
            ))
            cur_words, cur_text = [], ""
    if cur_words:
        segments.append((
            cur_words[0].get("start", 0),
            cur_words[-1].get("end", 0),
            cur_text,
        ))
    return segments


def _call_ai(prompt, ai_engine, api_key):
    if ai_engine == "claude":
        import anthropic
        model = _model_id(load_config().get("claude_model", "claude-sonnet-4-6"))
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model=model,
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text
    elif ai_engine == "ollama":
        import requests
        model = api_key.strip() if api_key.strip() else "qwen2.5:7b"
        resp = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": model, "prompt": prompt, "stream": False},
            timeout=300,
        )
        resp.raise_for_status()
        return resp.json()["response"]
    elif ai_engine == "lmstudio":
        import requests
        model = api_key.strip() if api_key.strip() else "google/gemma-4-e4b"
        resp = requests.post(
            "http://localhost:1234/v1/chat/completions",
            json={"model": model, "messages": [{"role": "user", "content": prompt}], "stream": False},
            timeout=300,
        )
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]
    elif ai_engine == "openai":
        from openai import OpenAI
        model = _model_id(load_config().get("openai_model", "gpt-4o"))
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.choices[0].message.content
    else:
        from google import genai
        model = _model_id(load_config().get("gemini_model", "gemini-3.1-flash-lite"))
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(
            model=model,
            contents=prompt,
        )
        return response.text


# 各 AI 引擎在 config.json 對應的 key 欄位，以及顯示名稱
_ENGINE_CFG_KEY = {
    "claude": "anthropic_key",
    "gemini": "gemini_key",
    "openai": "openai_key",
    "ollama": "ollama_model",
    "lmstudio": "lmstudio_model",
}
_ENGINE_DISPLAY = {
    "claude": "Claude",
    "gemini": "Gemini",
    "openai": "OpenAI",
    "ollama": "Ollama",
    "lmstudio": "LM Studio",
}

# ── 雲端引擎可選的模型清單（GUI 下拉選單用）────────────────────────────
# 維護方式：供應商新增/淘汰模型時，改下面這份清單再 git pull 即可，毋須改其他程式。
# 使用者在 GUI 仍可手動輸入未列出的模型名稱。本地引擎（ollama / lmstudio）的模型
# 直接打在 Key 欄位，不走這份清單。
_ENGINE_MODELS = {
    "claude": ["claude-opus-4-8", "claude-sonnet-4-6", "claude-haiku-4-5-20251001"],
    "gemini": [
        "gemini-3.1-pro-preview  ·  付費（品質最佳）",
        "gemini-2.5-pro  ·  付費",
        "gemini-3.5-flash  ·  免費",
        "gemini-2.5-flash  ·  免費",
        "gemini-3.1-flash-lite  ·  免費（最省）",
    ],
    "openai": ["gpt-4o", "gpt-4o-mini", "gpt-4.1"],
}

def _model_id(label):
    """從下拉顯示字串取回實際模型 ID（去掉「  ·  免費/付費」等標註）。
    使用者手動輸入純 ID 時不含分隔符，會原樣回傳。"""
    return (label or "").split("  ·  ")[0].strip()
# 各引擎的模型名稱在 config.json 對應的欄位（與 _call_ai 讀取的一致）
_ENGINE_MODEL_KEY = {
    "claude": "claude_model",
    "gemini": "gemini_model",
    "openai": "openai_model",
}
# 各引擎預設模型（與 _call_ai 的 fallback 一致；清單順序可不同，顯示以存檔值為準）
_ENGINE_MODEL_DEFAULT = {
    "claude": "claude-sonnet-4-6",
    "gemini": "gemini-3.1-flash-lite",
    "openai": "gpt-4o",
}

# ── 第一步驟「辨識引擎」清單（皆地端）─────────────────────────────────
# 顯示標籤 ←→ 內部值。語音辨識一律走地端（隱私、免費、且 WhisperX 能用聲紋
# 穩定分辨說話者）；雲端 AI 只在第二步驟做校正/摘要。
STT_ENGINES = [
    ("whisper",      "地端 Whisper"),
    ("whisperx",     "地端 WhisperX（含說話人辨識）"),
    ("whispercpp",   "地端 whisper.cpp（MacBook Air）"),
    ("whispercpp_diar", "地端 whisper.cpp ＋ 聲學分軌（whisperX，需 HF Token）"),
]
_STT_LABEL2VAL = {lb: v for v, lb in STT_ENGINES}
_STT_VAL2LABEL = {v: lb for v, lb in STT_ENGINES}

# Whisper 的 initial_prompt 約有 224 token 上限，超過會被截斷而失效；中文粗估
# 1 字 ≈ 1 token，這裡保守抓 180 字當整段預算，扣掉固定前綴後剩給「術語」的額度。
INITIAL_PROMPT_BUDGET = 180
_IP_BASE = "以下是繁體中文對話。"
_IP_LEAD = "可能提到："
_TERMS_BUDGET = INITIAL_PROMPT_BUDGET - len(_IP_BASE) - len(_IP_LEAD)


def _build_initial_prompt(terms):
    """把使用者輸入的關鍵術語組成 Whisper 的 initial_prompt，並裁在 token 預算內。
    超出的術語不會送進地端辨識（會走無限制的雲端/校正/摘要），避免整段被截掉而失效。"""
    terms = (terms or "").strip()
    if not terms:
        return _IP_BASE
    return (_IP_BASE + _IP_LEAD + terms)[:INITIAL_PROMPT_BUDGET]


def _total_ram_gb():
    """回傳系統實體記憶體（GB）；取不到回 None。"""
    try:
        import psutil
        return psutil.virtual_memory().total / (1024 ** 3)
    except Exception:
        pass
    try:
        if sys.platform == "darwin":
            out = subprocess.check_output(["sysctl", "-n", "hw.memsize"])
            return int(out.strip()) / (1024 ** 3)
    except Exception:
        pass
    try:
        return os.sysconf("SC_PAGE_SIZE") * os.sysconf("SC_PHYS_PAGES") / (1024 ** 3)
    except Exception:
        return None


def _suggest_model():
    """依記憶體建議地端預設模型：<10GB→small（如 8GB MBA），<24GB→medium，否則 large-v3。"""
    gb = _total_ram_gb()
    if gb is None:
        return "small"
    if gb < 10:
        return "small"
    if gb < 24:
        return "medium"
    return "large-v3"


def _suggest_stt_engine():
    """首次啟動時依機器條件建議地端引擎。MacBook Air 或記憶體 ≤ 8GB → whisper.cpp
    （輕、不吃記憶體、Mac 上穩）；其餘維持原本的 whisper。只在沒有存過引擎時呼叫。"""
    try:
        gb = _total_ram_gb()
        if gb is not None and gb <= 8.5:
            return "whispercpp"
        if sys.platform == "darwin":
            out = subprocess.run(["system_profiler", "SPHardwareDataType"],
                                 capture_output=True, text=True, timeout=8).stdout or ""
            if "MacBook Air" in out:
                return "whispercpp"
    except Exception:
        pass
    return "whisper"


def verify_engine(ai_engine, api_key):
    """以最小成本的呼叫驗證金鑰／連線，回傳 (ok: bool, message: str)。"""
    if ai_engine == "claude":
        import anthropic
        anthropic.Anthropic(api_key=api_key).models.list(limit=1)
        return True, "Claude 金鑰有效"
    elif ai_engine == "openai":
        from openai import OpenAI
        OpenAI(api_key=api_key).models.list()
        return True, "OpenAI 金鑰有效"
    elif ai_engine == "gemini":
        # 直接打 REST：SDK 在金鑰被拒時只丟出誤導性的「client has been closed」，
        # 改用 REST 才能拿到「API key not valid」等可讀的真正原因。
        import requests
        r = requests.get(
            "https://generativelanguage.googleapis.com/v1beta/models",
            params={"key": api_key}, timeout=15,
        )
        if r.status_code == 200:
            return True, "Gemini 金鑰有效"
        try:
            detail = r.json().get("error", {}).get("message", r.text[:200])
        except Exception:
            detail = r.text[:200]
        raise RuntimeError(f"HTTP {r.status_code}: {detail}")
    elif ai_engine == "ollama":
        import requests
        requests.get("http://localhost:11434/api/tags", timeout=10).raise_for_status()
        return True, "Ollama 連線正常"
    elif ai_engine == "lmstudio":
        import requests
        requests.get("http://localhost:1234/v1/models", timeout=10).raise_for_status()
        return True, "LM Studio 連線正常"
    return False, "未知引擎"


_TXT_TS_PAT = re.compile(r'^\[(\d+):(\d+)\s*-->\s*(\d+):(\d+)\]\s*(.+)$')
_TXT_SPK_PAT = re.compile(r'^(.{1,40})：\s*$')


def txt_to_srt(transcript_path, out_path=None):
    if out_path is None:
        base = transcript_path.rsplit(".", 1)[0]
        for suffix in ("_transcript_corrected", "_transcript"):
            if base.endswith(suffix):
                base = base[:-len(suffix)]
                break
        out_path = base + ".srt"

    def _fmt(total_seconds):
        h = total_seconds // 3600
        m = (total_seconds % 3600) // 60
        s = total_seconds % 60
        return f"{h:02d}:{m:02d}:{s:02d},000"

    entries = []
    current_speaker = None
    with open(transcript_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip()
            m = _TXT_TS_PAT.match(line)
            if m:
                sm, ss, em, es, text = m.groups()
                start = int(sm) * 60 + int(ss)
                end = int(em) * 60 + int(es)
                subtitle = text
                entries.append((_fmt(start), _fmt(end), subtitle))
            else:
                spk_m = _TXT_SPK_PAT.match(line)
                if spk_m and line.strip():
                    current_speaker = spk_m.group(1).strip()

    with open(out_path, "w", encoding="utf-8") as f:
        for i, (start, end, text) in enumerate(entries, 1):
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

    return out_path


_SRT_TS_PAT = re.compile(
    r'(\d+):(\d+):(\d+),(\d+)\s*-->\s*(\d+):(\d+):(\d+),(\d+)'
)


def _parse_srt(srt_path):
    """解析 SRT 檔，回傳 [(start_float, end_float, text), ...]。"""
    entries = []
    with open(srt_path, "r", encoding="utf-8") as f:
        content = f.read()
    for block in re.split(r'\n\s*\n', content.strip()):
        lines = block.strip().splitlines()
        for i, line in enumerate(lines):
            m = _SRT_TS_PAT.match(line.strip())
            if m:
                h1, m1, s1, ms1, h2, m2, s2, ms2 = m.groups()
                start = int(h1)*3600 + int(m1)*60 + int(s1) + int(ms1)/1000
                end   = int(h2)*3600 + int(m2)*60 + int(s2) + int(ms2)/1000
                text  = " ".join(lines[i+1:]).strip()
                if text:
                    entries.append((start, end, text))
                break
    return entries


def _ai_cut_segments(segments, ai_engine, api_key):
    """一次 API call 批次切割所有長 segment，回傳新的 [(start, end, text), ...]。"""
    short = {}   # index → (start, end, text) 不需要切的
    long_idx = []  # 需要切的 segment index

    for i, (start, end, text) in enumerate(segments):
        if len(text) <= 15:
            short[i] = (start, end, text)
        else:
            long_idx.append(i)

    phrases_map = {}  # index → [phrase, ...]
    if long_idx:
        seg_blocks = "\n".join(
            f"===SEG{i}===（{int(segments[i][1] - segments[i][0])}秒）\n{segments[i][2]}"
            for i in long_idx
        )
        prompt = AI_SRT_PROMPT.format(n=len(long_idx), segments=seg_blocks)
        ai_out = _call_ai(prompt, ai_engine, api_key)

        current_idx = None
        for line in ai_out.splitlines():
            line = line.strip()
            m = re.match(r'^===SEG(\d+)===', line)
            if m:
                current_idx = int(m.group(1))
                phrases_map[current_idx] = []
            elif current_idx is not None and line:
                phrases_map[current_idx].append(line)

    result = []
    for i, (start, end, text) in enumerate(segments):
        if i in short:
            result.append((start, end, text))
            continue
        duration = end - start
        phrases = phrases_map.get(i, [])
        if not phrases:
            result.append((start, end, text))
            continue
        total_chars = sum(len(p) for p in phrases)
        t = start
        for j, phrase in enumerate(phrases):
            ratio = len(phrase) / total_chars if total_chars > 0 else 1.0 / len(phrases)
            phrase_end = t + duration * ratio if j < len(phrases) - 1 else end
            result.append((t, phrase_end, phrase))
            t = phrase_end
    return result


def _write_srt(entries, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        for i, (s, e, text) in enumerate(entries, 1):
            f.write(f"{i}\n{_format_srt_time(s)} --> {_format_srt_time(e)}\n{text}\n\n")


def txt_to_srt_ai(transcript_path, ai_engine, api_key, out_path=None):
    segments = []
    with open(transcript_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip()
            m = _TXT_TS_PAT.match(line)
            if m:
                sm, ss, em, es, text = m.groups()
                segments.append((float(int(sm)*60 + int(ss)), float(int(em)*60 + int(es)), text.strip()))
    all_entries = _ai_cut_segments(segments, ai_engine, api_key)
    if out_path is None:
        base = transcript_path.rsplit(".", 1)[0]
        for suffix in ("_transcript_corrected", "_transcript"):
            if base.endswith(suffix):
                base = base[:-len(suffix)]
                break
        out_path = base + ".srt"
    _write_srt(all_entries, out_path)
    return out_path


def correct_transcript(transcript_path, ai_engine, api_key, out_path=None, context="", label_parties=False):
    """校正辨識錯誤。label_parties=True 時（稿子沒有說話人標記，如 whisper/whisper.cpp）
    同一次 API 順手依內容標出「哪一方」說的，省一半 token。"""
    with open(transcript_path, "r", encoding="utf-8") as f:
        transcript = f.read()
    ctx = f"\n背景資訊：{context}" if context else ""
    template = CORRECT_AND_PARTY_PROMPT if label_parties else CORRECT_PROMPT
    corrected = _call_ai(template.format(context=ctx, transcript=transcript), ai_engine, api_key)
    if out_path is None:
        base = transcript_path.rsplit("_transcript", 1)[0]
        out_path = base + "_transcript_corrected.txt"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(corrected)
    return out_path


def _open_file(path):
    """用系統預設程式開啟檔案（Mac / Windows / Linux 通用）。失敗不影響主流程。"""
    try:
        if sys.platform == "darwin":
            subprocess.run(["open", path], check=False)
        elif os.name == "nt":
            os.startfile(path)  # type: ignore[attr-defined]
        else:
            subprocess.run(["xdg-open", path], check=False)
    except Exception:
        pass


def _write_summary_docx(md_text, out_path):
    """把 AI 產出的 Markdown 摘要轉成精緻 .docx（標題分隔線、checkbox、清單、表格、**粗體**）。"""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def add_runs(paragraph, text):
        # 行內 **粗體**：split 後奇數段就是粗體內容
        for idx, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            if not part:
                continue
            run = paragraph.add_run(part)
            if idx % 2 == 1:
                run.bold = True

    def add_bottom_border(paragraph):
        # 在段落底部加一條淺灰分隔線（仿專業摘要的章節分隔）
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    doc = Document()
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        # 表格：連續的 | ... | 行
        if stripped.startswith("|") and stripped.endswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # 跳過分隔列 |---|---|
                if not all(c and set(c) <= set("-: ") for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=ncol)
                table.style = "Table Grid"
                for r in rows:
                    tcells = table.add_row().cells
                    for j in range(ncol):
                        tcells[j].text = r[j] if j < len(r) else ""
            continue
        # 標題（必須是 # 後接空白；# 人名 標籤無空白，不會誤判）
        hm = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            h = doc.add_heading(hm.group(2).strip(), level=level)
            h.paragraph_format.space_before = Pt(10 if level >= 2 else 0)
            h.paragraph_format.space_after = Pt(4)
            if level <= 2:  # H1 / H2 加章節分隔線
                add_bottom_border(h)
            i += 1
            continue
        # 行動項目 checkbox：- [ ] 或 - [x]
        cm = re.match(r"^[-*]\s+\[([ xX])\]\s+(.*)", stripped)
        if cm:
            box = "☑" if cm.group(1).lower() == "x" else "☐"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.add_run(box + "  ")
            add_runs(p, cm.group(2).strip())
            i += 1
            continue
        # 項目清單
        if stripped.startswith(("- ", "* ")):
            add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:].strip())
            i += 1
            continue
        nm = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if nm:
            add_runs(doc.add_paragraph(style="List Number"), nm.group(2).strip())
            i += 1
            continue
        # 一般段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, stripped)
        i += 1

    doc.save(out_path)


def _clean_basename(path):
    """去掉 _transcript / _transcript_corrected / 副檔名，還原成錄音檔名當標題。"""
    base = os.path.basename(path)
    for suf in ("_transcript_corrected.txt", "_transcript.txt", ".txt"):
        if base.endswith(suf):
            return base[:-len(suf)]
    return os.path.splitext(base)[0]


def _guess_meeting_time(path):
    """推測會議時間：先試從檔名抓，抓不到就用檔案修改時間。"""
    name = os.path.basename(path)
    mtime = datetime.fromtimestamp(os.path.getmtime(path))
    # 檔名含 "6月10日 14-56" / "6月10日 14:56"
    m = re.search(r'(\d{1,2})月(\d{1,2})日\s*(\d{1,2})[-:](\d{2})', name)
    if m:
        mon, day, hh, mm = (int(x) for x in m.groups())
        return f"{mtime.year}-{mon:02d}-{day:02d} {hh:02d}:{mm:02d}"
    # 檔名含 "2026-06-10 14-56" / "2026-06-10_1456"
    m = re.search(r'(\d{4})-(\d{1,2})-(\d{1,2})[ _]?(\d{2})[-:]?(\d{2})', name)
    if m:
        y, mo, d, hh, mm = (int(x) for x in m.groups())
        return f"{y}-{mo:02d}-{d:02d} {hh:02d}:{mm:02d}"
    # 抓不到 → 用檔案時間
    return mtime.strftime("%Y-%m-%d %H:%M")


def generate_summary(transcript_path, ai_engine, api_key, out_path=None, context=""):
    with open(transcript_path, "r", encoding="utf-8") as f:
        transcript = f.read()
    ctx = f"\n背景資訊：{context}" if context else ""
    title = _clean_basename(transcript_path)            # 標題＝錄音檔名
    meeting_time = _guess_meeting_time(transcript_path)  # 會議時間（檔名或檔案時間）
    notes = _call_ai(
        NOTES_PROMPT.format(context=ctx, transcript=transcript, meeting_time=meeting_time),
        ai_engine, api_key)
    md = f"# {title}\n\n{notes}"                          # 標題由系統以檔名填入，不靠 AI 生
    if out_path is None:
        out_path = os.path.join(os.path.dirname(transcript_path), f"{title}_摘要.docx")
    _write_summary_docx(md, out_path)
    return out_path


class _StdoutCapture:
    _pat = re.compile(r'\[(\d+):(\d+\.\d+)')

    def __init__(self, duration, prog_q):
        self._duration = duration
        self._prog_q = prog_q
        self._buf = ""

    def write(self, text):
        self._buf += text
        while '\n' in self._buf:
            line, self._buf = self._buf.split('\n', 1)
            m = self._pat.search(line)
            if m and self._duration > 0:
                t = int(m.group(1)) * 60 + float(m.group(2))
                pct = min(int(t / self._duration * 100), 99)
                self._prog_q.put(pct)

    def flush(self):
        pass


class _ScaledProg:
    """把「單一切段內 0–100% 的進度」換算成「整檔（第 i/n 段）的全域百分比」再轉發，
    讓切塊處理時進度條在每一段裡面也會平滑前進，而不是一段跳一格。"""

    def __init__(self, real_q, i, n):
        self._q, self._i, self._n = real_q, i, n

    def put(self, p):
        frac = min(max(p, 0), 100) / 100.0
        self._q.put(int((self._i + frac) / self._n * 100))


def _collapse_repeats(text):
    """壓掉 Whisper 幻覺式重複（對對對對…、我懂了 我懂了…），保留正常內容。"""
    if not text:
        return text
    # 逗號/頓號分隔的重複：對,對,對,對
    text = re.sub(r'([一-鿿！？。]{1,8})([,，、]\1){2,}', r'\1', text)
    # 無分隔短字重複：對對對對 / 偏偏偏偏（連 4 次以上才壓，保留「謝謝」「哈哈哈」）
    text = re.sub(r'([一-鿿]{1,4})\1{3,}', r'\1', text)
    # 空白分隔短句重複：我懂了 我懂了 我懂了
    text = re.sub(r'([一-鿿，。？！、]{2,10})( \1){2,}', r'\1', text)
    # 完整句子重複（以標點結尾的長句）
    text = re.sub(r'([一-鿿，。？！、 ]{8,30}[？。，！])\s*(\1\s*){1,}', r'\1', text)
    return text


# ── 長檔切段 + checkpoint ──────────────────────────────────────────────
# 動機：長音檔在地端轉錄/說話人辨識會吃大量記憶體（pyannote 吃整檔），且「跑到
# 一半崩潰就全盤皆輸」。超過門檻的檔案會先在「靜音處」就近切成數段，逐段轉錄，
# 每段一完成就寫一份 checkpoint；任何一段失敗，前面成功的段落都保留並縫成「部分
# 逐字稿」。切段會讓 pyannote 的說話人編號跨段不一致，但這由第二步驟的 AI 校正
# 依內容重新統一「哪一方」，故此處不強求跨段聲紋一致。
CHUNK_THRESHOLD_SEC = 1200   # 超過 20 分鐘才啟用切段（短檔走原本單次流程，零影響）
CHUNK_TARGET_SEC = 600       # 每段目標長度約 10 分鐘，於最接近的靜音處下刀


def _json_default(o):
    try:
        return float(o)
    except Exception:
        return str(o)


def _detect_silences(audio_path):
    """用 ffmpeg silencedetect 找靜音區間，回傳 [(start, end), ...]（秒）。"""
    import subprocess
    cmd = ["ffmpeg", "-hide_banner", "-vn", "-i", audio_path,
           "-af", "silencedetect=noise=-30dB:d=0.5", "-f", "null", "-"]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True)
    except Exception:
        return []
    err = proc.stderr or ""
    starts = [float(x) for x in re.findall(r"silence_start:\s*([0-9.]+)", err)]
    ends = [float(x) for x in re.findall(r"silence_end:\s*([0-9.]+)", err)]
    sil = []
    for i, e in enumerate(ends):
        s = starts[i] if i < len(starts) else max(0.0, e - 0.5)
        sil.append((s, e))
    return sil


def _plan_cuts(duration, silences, target=CHUNK_TARGET_SEC):
    """規劃切點：每隔約 target 秒，挑最接近的靜音中點下刀；找不到靜音才硬切。
    回傳 [(start, end), ...] 涵蓋整個檔案。"""
    mids = sorted((s + e) / 2 for s, e in silences)
    cuts, last = [], 0.0
    while duration - last > target * 1.5:
        ideal = last + target
        lo, hi = last + target * 0.5, last + target * 1.5
        cand = [m for m in mids if lo <= m <= hi]
        cut = min(cand, key=lambda m: abs(m - ideal)) if cand else ideal
        cuts.append(round(cut, 3))
        last = cut
    bounds, prev = [], 0.0
    for c in cuts:
        bounds.append((prev, c))
        prev = c
    bounds.append((prev, round(duration, 3)))
    return bounds


def _cut_audio(audio_path, bounds, work_dir):
    """依 bounds 切出 16k 單聲道 wav；已存在的段沿用（支援續傳）。
    回傳 [(chunk_path, start_offset_sec), ...]。"""
    import subprocess
    paths = []
    for i, (s, e) in enumerate(bounds):
        out = os.path.join(work_dir, f"chunk_{i:03d}.wav")
        if not os.path.exists(out):
            cmd = ["ffmpeg", "-hide_banner", "-y", "-ss", f"{s:.3f}", "-i", audio_path,
                   "-t", f"{e - s:.3f}", "-ar", "16000", "-ac", "1", out]
            subprocess.run(cmd, capture_output=True)
        paths.append((out, s))
    return paths


def _offset_segments(segs, off):
    """把一段內所有時間戳（含逐字 words）加上該段的起始偏移，回到全片時間軸。"""
    out = []
    for s in segs:
        s = dict(s)
        if s.get("start") is not None:
            s["start"] = float(s["start"]) + off
        if s.get("end") is not None:
            s["end"] = float(s["end"]) + off
        if s.get("words"):
            ws = []
            for w in s["words"]:
                w = dict(w)
                if w.get("start") is not None:
                    w["start"] = float(w["start"]) + off
                if w.get("end") is not None:
                    w["end"] = float(w["end"]) + off
                ws.append(w)
            s["words"] = ws
        out.append(s)
    return out


def _plan_or_load_bounds(audio, duration, work_dir, log_q):
    """讀既有 manifest（續傳時切點一致）或重新規劃切點並寫 manifest。"""
    import json
    manifest = os.path.join(work_dir, "manifest.json")
    if os.path.exists(manifest):
        try:
            with open(manifest, encoding="utf-8") as f:
                return [tuple(b) for b in json.load(f)["bounds"]]
        except Exception:
            pass
    log_q.put("分析靜音以決定切點...")
    bounds = _plan_cuts(duration, _detect_silences(audio))
    with open(manifest, "w", encoding="utf-8") as f:
        json.dump({"bounds": bounds}, f)
    return bounds


def _collect_done_chunks(work_dir, n):
    """收集 0..n-1 連續已完成的 chunk json；遇到缺號即停（避免時間軸出現空洞）。"""
    import json
    all_segs, done = [], 0
    for i in range(n):
        cj = os.path.join(work_dir, f"chunk_{i:03d}.json")
        if not os.path.exists(cj):
            break
        with open(cj, encoding="utf-8") as f:
            all_segs.extend(json.load(f))
        done += 1
    return all_segs, done


def _write_transcript_plain(out_path, all_segs):
    if os.path.exists(out_path):
        os.remove(out_path)  # 新檔的建立/加入日期=現在，免得 Finder 依加入日期排序排回舊日期
    with open(out_path, "w", encoding="utf-8") as f:
        for seg in all_segs:
            start, end = int(seg["start"]), int(seg["end"])
            ts = f"[{start//60:02d}:{start%60:02d} --> {end//60:02d}:{end%60:02d}]"
            f.write(f"{ts} {_collapse_repeats(seg['text'].strip())}\n")


def _write_transcript_speaker(out_path, all_segs, converter):
    if os.path.exists(out_path):
        os.remove(out_path)  # 同上：新檔日期=現在
    with open(out_path, "w", encoding="utf-8") as f:
        current = None
        for seg in all_segs:
            start, end = int(seg.get("start", 0)), int(seg.get("end", 0))
            speaker = seg.get("speaker", "SPEAKER_??")
            text = seg.get("text", "").strip()
            if not text:
                continue
            if converter:
                text = converter.convert(text)
            text = _collapse_repeats(text)
            ts = f"[{start//60:02d}:{start%60:02d} --> {end//60:02d}:{end%60:02d}]"
            if speaker != current:
                if current is not None:
                    f.write("\n")
                f.write(f"{speaker}：\n")
                current = speaker
            f.write(f"{ts} {text}\n")


def _write_srt_segments(srt_path, all_segs, converter=None):
    all_chunks = []
    for seg in all_segs:
        words = seg.get("words", [])
        if words:
            for s, e, text in _words_to_srt_segments(words):
                if converter:
                    text = converter.convert(text)
                all_chunks.append((s, e, text))
        else:
            text = seg.get("text", "").strip()
            if converter:
                text = converter.convert(text)
            all_chunks.append((seg.get("start", 0), seg.get("end", 0), text))
    with open(srt_path, "w", encoding="utf-8") as f:
        for i, (s, e, text) in enumerate(all_chunks, 1):
            f.write(f"{i}\n{_format_srt_time(s)} --> {_format_srt_time(e)}\n{text}\n\n")


def _finalize_chunks(audio, out_dir, work_dir, n, failed, write_srt,
                     speaker_mode, converter, result_q, log_q):
    """縫合已完成的段落；全完成→寫正式稿並清掉工作目錄，部分完成→存『部分逐字稿』並保留工作目錄供續傳。"""
    import shutil
    all_segs, done = _collect_done_chunks(work_dir, n)
    if done == 0:
        result_q.put(("error", failed or "所有段落皆轉錄失敗"))
        return
    base = os.path.splitext(os.path.basename(audio))[0]
    partial = done < n
    suffix = "_transcript_部分完成.txt" if partial else "_transcript.txt"
    out_path = os.path.join(out_dir, f"{base}{suffix}")
    if speaker_mode:
        _write_transcript_speaker(out_path, all_segs, converter)
    else:
        _write_transcript_plain(out_path, all_segs)
    if write_srt:
        srt_path = os.path.join(out_dir, f"{base}{'_部分' if partial else ''}.srt")
        _write_srt_segments(srt_path, all_segs, converter)
        log_q.put(f"字幕已輸出：{os.path.basename(srt_path)}")
    if partial:
        log_q.put(f"⚠ 只完成 {done}/{n} 段（最後一段失敗），已保留並存成「{os.path.basename(out_path)}」。"
                  f"修正問題後重新轉錄同一檔，會自動從第 {done + 1} 段續跑。")
    else:
        shutil.rmtree(work_dir, ignore_errors=True)
        # 清掉先前失敗留下的「部分完成」殘檔，避免混淆
        for stale in (os.path.join(out_dir, f"{base}_transcript_部分完成.txt"),
                      os.path.join(out_dir, f"{base}_部分.srt")):
            try:
                os.remove(stale)
            except OSError:
                pass
        log_q.put(f"全部 {n} 段完成並縫合。")
    result_q.put(("done", out_path))


def _run_whisper_chunked(model, audio, out_dir, lang, prompt, write_srt,
                         duration, result_q, log_q, prog_q):
    import json
    base = os.path.splitext(os.path.basename(audio))[0]
    work_dir = os.path.join(out_dir, f".voxlog_chunks_{base}")
    os.makedirs(work_dir, exist_ok=True)
    bounds = _plan_or_load_bounds(audio, duration, work_dir, log_q)
    chunks = _cut_audio(audio, bounds, work_dir)
    n = len(chunks)
    log_q.put(f"長檔（{int(duration//60)} 分）已切成 {n} 段，逐段轉錄並建立 checkpoint...")
    failed = None
    for i, (cpath, off) in enumerate(chunks):
        cj = os.path.join(work_dir, f"chunk_{i:03d}.json")
        if os.path.exists(cj):
            log_q.put(f"段 {i+1}/{n} 已有 checkpoint，沿用。")
        else:
            try:
                log_q.put(f"轉錄段 {i+1}/{n} ...")
                # 用 stdout 擷取 whisper 的逐段時間，換算成全域進度 → 段內也會平滑前進
                chunk_dur = bounds[i][1] - bounds[i][0]
                _saved_stdout = sys.stdout
                sys.stdout = _StdoutCapture(chunk_dur, _ScaledProg(prog_q, i, n))
                try:
                    result = model.transcribe(cpath, language=lang,
                                              initial_prompt=prompt or None,
                                              word_timestamps=True, verbose=True)
                finally:
                    sys.stdout = _saved_stdout
                segs = _offset_segments(result["segments"], off)
                with open(cj, "w", encoding="utf-8") as f:
                    json.dump(segs, f, ensure_ascii=False, default=_json_default)
            except Exception as e:
                failed = str(e)
                log_q.put(f"⚠ 段 {i+1}/{n} 失敗：{e}")
                break
        prog_q.put(int((i + 1) / n * 100))
    _finalize_chunks(audio, out_dir, work_dir, n, failed, write_srt,
                     speaker_mode=False, converter=None,
                     result_q=result_q, log_q=log_q)


def _run_whisperx_chunked(model, audio, out_dir, lang, hf_token, num_speakers,
                          write_srt, device, duration, result_q, log_q, prog_q):
    import json, whisperx
    from whisperx.diarize import DiarizationPipeline
    try:
        import opencc
        converter = opencc.OpenCC("s2twp")
    except Exception:
        converter = None
    base = os.path.splitext(os.path.basename(audio))[0]
    work_dir = os.path.join(out_dir, f".voxlog_chunks_{base}")
    os.makedirs(work_dir, exist_ok=True)
    bounds = _plan_or_load_bounds(audio, duration, work_dir, log_q)
    chunks = _cut_audio(audio, bounds, work_dir)
    n = len(chunks)
    log_q.put(f"長檔（{int(duration//60)} 分）已切成 {n} 段，逐段轉錄＋說話人辨識並建立 checkpoint...")
    log_q.put("提醒：切段後說話人編號跨段不保證一致，請於第二步驟用 AI 校正依內容統一『哪一方』。")
    diarize_model = DiarizationPipeline(token=hf_token, device=device)
    model_a = metadata = align_lang = None
    failed = None
    for i, (cpath, off) in enumerate(chunks):
        cj = os.path.join(work_dir, f"chunk_{i:03d}.json")
        if os.path.exists(cj):
            log_q.put(f"段 {i+1}/{n} 已有 checkpoint，沿用。")
            prog_q.put(int((i + 1) / n * 100))
            continue
        try:
            log_q.put(f"轉錄段 {i+1}/{n} ...")
            audio_data = whisperx.load_audio(cpath)
            tkw = {"batch_size": 8}
            if lang:
                tkw["language"] = lang
            result = model.transcribe(audio_data, **tkw)
            prog_q.put(int((i + 0.4) / n * 100))   # 轉錄完
            dl = result.get("language", lang or "zh")
            if model_a is None or dl != align_lang:
                model_a, metadata = whisperx.load_align_model(language_code=dl, device=device)
                align_lang = dl
            log_q.put(f"段 {i+1}/{n}：對齊時間戳...")
            result = whisperx.align(result["segments"], model_a, metadata, audio_data,
                                    device, return_char_alignments=False)
            prog_q.put(int((i + 0.65) / n * 100))  # 對齊完
            log_q.put(f"段 {i+1}/{n}：分析說話人...")
            dkw = {}
            if num_speakers and num_speakers > 0:
                dkw["num_speakers"] = num_speakers
            diarize_segments = diarize_model(audio_data, **dkw)
            result = whisperx.assign_word_speakers(diarize_segments, result)
            prog_q.put(int((i + 0.9) / n * 100))   # 說話人辨識完
            segs = _offset_segments(result["segments"], off)
            with open(cj, "w", encoding="utf-8") as f:
                json.dump(segs, f, ensure_ascii=False, default=_json_default)
            del audio_data
        except Exception as e:
            failed = str(e)
            log_q.put(f"⚠ 段 {i+1}/{n} 失敗：{e}")
            break
        prog_q.put(int((i + 1) / n * 100))
    _finalize_chunks(audio, out_dir, work_dir, n, failed, write_srt,
                     speaker_mode=True, converter=converter,
                     result_q=result_q, log_q=log_q)


def _safe_open_transcript(out_path, log_q):
    """寫逐字稿前統一處理：先刪同名舊檔再建新檔（macOS 上「刪舊+建新」TCC 允許，
    可避開「原地截斷覆寫」在 Downloads/Desktop/Documents 被擋的 EPERM；也讓 Finder
    依現在日期排序）。主路徑被系統擋下時 fallback 到 ~/VoxLog。回傳 (file_obj, 實際路徑)。"""
    import os
    candidates = [out_path]
    fb_path = os.path.join(os.path.expanduser("~"), "VoxLog", os.path.basename(out_path))
    if os.path.abspath(fb_path) != os.path.abspath(out_path):
        candidates.append(fb_path)
    last_err = None
    for cand in candidates:
        try:
            os.makedirs(os.path.dirname(cand), exist_ok=True)
            if os.path.exists(cand):
                os.remove(cand)
            f = open(cand, "w", encoding="utf-8")
            if cand != out_path:
                log_q.put(f"⚠ 「{os.path.dirname(out_path)}」寫入被系統擋下（隱私保護），已改存到：{cand}")
            return f, cand
        except (PermissionError, OSError) as e:
            last_err = e
    raise last_err


def whisper_worker(audio, out_dir, model_name, lang, prompt, write_srt, result_q, log_q, prog_q, out_path_override=None):
    try:
        import whisper, torch, os, sys, platform
        if platform.system() == "Windows":
            os.environ["PATH"] += ";" + r"C:\Users\kamiy\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.1.1-full_build\bin"

        if torch.cuda.is_available():
            device, gpu_name = "cuda", torch.cuda.get_device_name(0)
        elif hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
            # Apple MPS 缺少 Whisper 用到的稀疏張量運算
            # (aten::_sparse_coo_tensor_with_dims_and_tensors / SparseMPS backend)，
            # 在 MPS 上轉錄會直接崩潰，因此 Mac 上的一般 Whisper 引擎改用 CPU。
            device, gpu_name = "cpu", "CPU（Mac 不支援 MPS 轉錄，已自動改用 CPU）"
        else:
            device, gpu_name = "cpu", "無"

        log_q.put(f"使用裝置：{device.upper()}（{gpu_name}）")
        log_q.put(f"載入模型：{model_name}")
        model = whisper.load_model(model_name, device=device)

        audio_data = whisper.load_audio(audio)
        duration = len(audio_data) / 16000
        m, s = int(duration // 60), int(duration % 60)
        log_q.put(f"開始轉錄：{os.path.basename(audio)}（{m} 分 {s} 秒）")

        if duration > CHUNK_THRESHOLD_SEC:
            del audio_data  # 改走切段流程，不需整檔常駐記憶體
            _run_whisper_chunked(model, audio, out_dir, lang, prompt, write_srt,
                                 duration, result_q, log_q, prog_q)
            return

        sys.stdout = _StdoutCapture(duration, prog_q)
        result = model.transcribe(audio, language=lang, verbose=True,
                                  initial_prompt=prompt or None,
                                  word_timestamps=True)
        sys.stdout = sys.__stdout__
        prog_q.put(100)

        base = os.path.splitext(os.path.basename(audio))[0]
        out_path = out_path_override or os.path.join(out_dir, f"{base}_transcript.txt")
        f, out_path = _safe_open_transcript(out_path, log_q)
        with f:
            for seg in result["segments"]:
                start = int(seg["start"])
                end = int(seg["end"])
                ts = f"[{start//60:02d}:{start%60:02d} --> {end//60:02d}:{end%60:02d}]"
                f.write(f"{ts} {_collapse_repeats(seg['text'].strip())}\n")

        if write_srt:
            srt_path = os.path.join(out_dir, f"{base}.srt")
            all_chunks = []
            for seg in result["segments"]:
                words = seg.get("words", [])
                if words:
                    all_chunks.extend(_words_to_srt_segments(words))
                else:
                    all_chunks.append((seg["start"], seg["end"], seg["text"].strip()))
            with open(srt_path, "w", encoding="utf-8") as f:
                for i, (s, e, text) in enumerate(all_chunks, 1):
                    f.write(f"{i}\n{_format_srt_time(s)} --> {_format_srt_time(e)}\n{text}\n\n")
            log_q.put(f"字幕已輸出：{os.path.basename(srt_path)}")

        result_q.put(("done", out_path))
    except Exception as e:
        try:
            sys.stdout = sys.__stdout__
        except Exception:
            pass
        result_q.put(("error", str(e)))


def whisperx_worker(audio, out_dir, model_name, lang, hf_token, num_speakers, prompt, write_srt, result_q, log_q, prog_q, out_path_override=None):
    try:
        import whisperx, torch, os, platform
        if platform.system() == "Windows":
            os.environ["PATH"] += ";" + r"C:\Users\kamiy\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.1.1-full_build\bin"

        if torch.cuda.is_available():
            device, gpu_name = "cuda", torch.cuda.get_device_name(0)
            # 舊卡（Pascal/Maxwell 等，compute capability < 7.0）FP16 很弱，
            # 用 int8_float16 反而慢/不穩，自動降成純 int8。Volta(7.0)+ 才用 int8_float16。
            try:
                cc_major = torch.cuda.get_device_capability(0)[0]
            except Exception:
                cc_major = 7
            if cc_major >= 7:
                compute_type = "int8_float16"
            else:
                compute_type = "int8"
                log_q.put(f"偵測到較舊的顯卡（{gpu_name}），自動改用 int8 計算（FP16 在這代卡較慢）。")
        else:
            device, compute_type, gpu_name = "cpu", "int8", "無（CPU）"

        log_q.put(f"使用裝置：{device.upper()}（{gpu_name}）")
        log_q.put(f"載入模型：{model_name}")

        model = whisperx.load_model(model_name, device, compute_type=compute_type)
        prog_q.put(10)

        audio_data = whisperx.load_audio(audio)
        duration = len(audio_data) / 16000
        m, s = int(duration // 60), int(duration % 60)
        log_q.put(f"開始轉錄：{os.path.basename(audio)}（{m} 分 {s} 秒）")
        if duration > CHUNK_THRESHOLD_SEC:
            # 長檔：切段 + checkpoint，避免 pyannote 吃整檔記憶體與「崩潰全盤皆輸」
            del audio_data
            _run_whisperx_chunked(model, audio, out_dir, lang, hf_token, num_speakers,
                                  write_srt, device, duration, result_q, log_q, prog_q)
            return

        transcribe_kwargs = {"batch_size": 8}
        if lang:
            transcribe_kwargs["language"] = lang
        result = model.transcribe(audio_data, **transcribe_kwargs)
        prog_q.put(40)

        detected_lang = result.get("language", lang or "zh")
        log_q.put("對齊時間戳...")
        model_a, metadata = whisperx.load_align_model(language_code=detected_lang, device=device)
        result = whisperx.align(result["segments"], model_a, metadata, audio_data, device,
                                return_char_alignments=False)
        prog_q.put(65)

        log_q.put("分析說話人...")
        from whisperx.diarize import DiarizationPipeline
        diarize_model = DiarizationPipeline(token=hf_token, device=device)
        diarize_kwargs = {}
        if num_speakers and num_speakers > 0:
            diarize_kwargs["num_speakers"] = num_speakers
        diarize_segments = diarize_model(audio_data, **diarize_kwargs)
        result = whisperx.assign_word_speakers(diarize_segments, result)
        prog_q.put(90)

        try:
            import opencc
            converter = opencc.OpenCC("s2twp")
        except Exception:
            converter = None

        base = os.path.splitext(os.path.basename(audio))[0]
        out_path = out_path_override or os.path.join(out_dir, f"{base}_transcript.txt")
        segments_for_srt = []
        f, out_path = _safe_open_transcript(out_path, log_q)
        with f:
            current_speaker = None
            for seg in result["segments"]:
                start = seg.get("start", 0)
                end = seg.get("end", 0)
                speaker = seg.get("speaker", "SPEAKER_??")
                text = seg.get("text", "").strip()
                if not text:
                    continue
                if converter:
                    text = converter.convert(text)
                text = _collapse_repeats(text)
                ts = f"[{int(start)//60:02d}:{int(start)%60:02d} --> {int(end)//60:02d}:{int(end)%60:02d}]"
                if speaker != current_speaker:
                    if current_speaker is not None:
                        f.write("\n")
                    f.write(f"{speaker}：\n")
                    current_speaker = speaker
                f.write(f"{ts} {text}\n")
                segments_for_srt.append((start, end, speaker, text))

        if write_srt:
            srt_path = os.path.join(out_dir, f"{base}.srt")
            all_chunks = []
            for seg in result["segments"]:
                speaker = seg.get("speaker", "")
                words = seg.get("words", [])
                if words:
                    for s, e, text in _words_to_srt_segments(words):
                        if converter:
                            text = converter.convert(text)
                        all_chunks.append((s, e, speaker, text))
                else:
                    text = seg.get("text", "").strip()
                    if converter:
                        text = converter.convert(text)
                    all_chunks.append((seg.get("start", 0), seg.get("end", 0), speaker, text))
            with open(srt_path, "w", encoding="utf-8") as f:
                for i, (s, e, speaker, text) in enumerate(all_chunks, 1):
                    f.write(f"{i}\n{_format_srt_time(s)} --> {_format_srt_time(e)}\n{text}\n\n")
            log_q.put(f"字幕已輸出：{os.path.basename(srt_path)}")

        result_q.put(("done", out_path))
    except Exception as e:
        result_q.put(("error", str(e)))


# ── whisper.cpp 引擎（輕量、低記憶體，給 MacBook Air 等機器）─────────────
# 走外部 whisper-cli（brew install whisper-cpp / Windows 預編譯）+ ggml 量化模型，
# 記憶體佔用極低且不隨音檔長度成長。無內建說話人辨識——「哪一方」交給第二步驟
# 的 AI 校正依內容判斷。模型首次使用會自動下載到專案的 models/ 目錄。
_WHISPERCPP_MODELS = {
    "tiny": "ggml-tiny.bin",
    "base": "ggml-base.bin",
    "small": "ggml-small-q5_1.bin",      # 量化版，省記憶體
    "medium": "ggml-medium-q5_0.bin",    # MacBook Air 建議用這個（準度/記憶體最佳平衡）
    "large": "ggml-large-v3.bin",
    "large-v3": "ggml-large-v3.bin",
}
_WHISPERCPP_BASE_URL = "https://huggingface.co/ggerganov/whisper.cpp/resolve/main/"


def _resolve_whispercpp_binary():
    """找 whisper-cli。優先用 repo 內附帶的 binary（Windows 同事 git pull 即可用、免裝免設 PATH），
    找不到才退回系統 PATH。回傳路徑字串，或 None。"""
    here = os.path.dirname(os.path.abspath(__file__))
    if platform.system() == "Windows":
        bundled = [os.path.join(here, "bin", "whispercpp-win", "whisper-cli.exe"),
                   os.path.join(here, "bin", "whispercpp-win", "main.exe")]
    else:
        bundled = [os.path.join(here, "bin", "whispercpp-mac", "whisper-cli")]
    for b in bundled:
        if os.path.exists(b):
            # 把 binary 所在資料夾加進 PATH，確保它旁邊的 dll/dylib 找得到
            d = os.path.dirname(b)
            if d not in os.environ.get("PATH", ""):
                os.environ["PATH"] = d + os.pathsep + os.environ.get("PATH", "")
            return b
    return shutil.which("whisper-cli") or shutil.which("whisper-cpp")


def _ensure_whispercpp_model(size, models_dir, log_q):
    """確保 ggml 模型檔存在，不存在則自動下載。回傳模型路徑。"""
    import requests
    fname = _WHISPERCPP_MODELS.get(size, f"ggml-{size}.bin")
    path = os.path.join(models_dir, fname)
    if os.path.exists(path):
        return path
    os.makedirs(models_dir, exist_ok=True)
    log_q.put(f"首次使用 whisper.cpp「{size}」模型，下載中（{fname}）…首次需一兩分鐘，之後就不必再等。")
    tmp = path + ".part"
    # 用 requests（自帶 certifi 憑證）串流下載：macOS 上 urllib 找不到系統 CA 憑證，
    # 會報 CERTIFICATE_VERIFY_FAILED；串流寫檔也不會把整包大模型讀進記憶體。
    with requests.get(_WHISPERCPP_BASE_URL + fname, stream=True, timeout=60) as r:
        r.raise_for_status()
        with open(tmp, "wb") as f:
            for chunk in r.iter_content(chunk_size=1 << 20):
                if chunk:
                    f.write(chunk)
    os.replace(tmp, path)
    log_q.put(f"模型下載完成：{fname}")
    return path


def whispercpp_worker(audio, out_dir, model_name, lang, prompt, write_srt, result_q, log_q, prog_q, out_path_override=None):
    import subprocess, shutil, json, tempfile, platform, re as _re
    try:
        if platform.system() == "Windows":
            os.environ["PATH"] += ";" + r"C:\Users\kamiy\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.1.1-full_build\bin"

        binary = _resolve_whispercpp_binary()
        if not binary:
            result_q.put(("error",
                "找不到 whisper-cli。\nMac 請在終端機執行：brew install whisper-cpp\n"
                "Windows 請安裝 whisper.cpp 並將其加入系統 PATH（或放進 bin/whispercpp-win/）。"))
            return

        models_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models")
        log_q.put(f"使用引擎：whisper.cpp（模型 {model_name}）")
        model_path = _ensure_whispercpp_model(model_name, models_dir, log_q)

        # whisper.cpp 需要 16k 單聲道 wav，先用 ffmpeg 轉檔
        log_q.put(f"準備音檔：{os.path.basename(audio)}")
        wav = os.path.join(tempfile.gettempdir(), "voxlog_wcpp_input.wav")
        subprocess.run(["ffmpeg", "-hide_banner", "-loglevel", "error", "-y", "-i", audio,
                        "-ar", "16000", "-ac", "1", wav], check=True)

        base = os.path.splitext(os.path.basename(audio))[0]
        out_json_base = os.path.join(tempfile.gettempdir(), f"voxlog_wcpp_{base}")
        cmd = [binary, "-m", model_path, "-f", wav, "-oj", "-of", out_json_base, "-pp",
               "-l", lang or "auto"]
        if prompt:
            cmd += ["--prompt", prompt]
        log_q.put("whisper.cpp 轉錄中…")
        proc = subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE, text=True)
        for line in proc.stderr:
            m = _re.search(r"progress\s*=\s*(\d+)%", line)
            if m:
                prog_q.put(int(m.group(1)))
        proc.wait()
        if proc.returncode != 0:
            result_q.put(("error", f"whisper.cpp 轉錄失敗（return code {proc.returncode}）"))
            return
        prog_q.put(100)

        with open(out_json_base + ".json", encoding="utf-8") as f:
            segs = json.load(f).get("transcription", [])

        try:
            import opencc
            converter = opencc.OpenCC("s2twp")
        except Exception:
            converter = None

        out_path = out_path_override or os.path.join(out_dir, f"{base}_transcript.txt")
        srt_entries = []
        f, out_path = _safe_open_transcript(out_path, log_q)
        with f:
            for seg in segs:
                off = seg.get("offsets", {})
                start = (off.get("from", 0) or 0) / 1000.0
                end = (off.get("to", 0) or 0) / 1000.0
                text = (seg.get("text") or "").strip()
                if not text:
                    continue
                if converter:
                    text = converter.convert(text)
                text = _collapse_repeats(text)
                ts = f"[{int(start)//60:02d}:{int(start)%60:02d} --> {int(end)//60:02d}:{int(end)%60:02d}]"
                f.write(f"{ts} {text}\n")
                srt_entries.append((start, end, text))

        if write_srt:
            srt_path = os.path.join(out_dir, f"{base}.srt")
            with open(srt_path, "w", encoding="utf-8") as f:
                for i, (s, e, text) in enumerate(srt_entries, 1):
                    f.write(f"{i}\n{_format_srt_time(s)} --> {_format_srt_time(e)}\n{text}\n\n")
            log_q.put(f"字幕已輸出：{os.path.basename(srt_path)}")

        log_q.put("註：whisper.cpp 沒有說話人標記；可在第二步驟用 AI 校正依內容標出「哪一方」。")
        for tmp_f in (wav, out_json_base + ".json"):
            try:
                os.remove(tmp_f)
            except OSError:
                pass
        result_q.put(("done", out_path))
    except Exception as e:
        result_q.put(("error", str(e)))


def whispercpp_diar_worker(audio, out_dir, model_name, lang, hf_token, num_speakers, prompt, write_srt, result_q, log_q, prog_q, out_path_override=None):
    """混搭引擎：① whisper.cpp（Metal，快）負責轉錄 → ② whisperX align + pyannote 分軌
    負責說話者。接續執行（非平行），峰值負載=較重的單一階段。輸出帶 SPEAKER_xx 的稿，
    之後可在第二步用「修正說話者」改成真實角色名。進度：0–65% 轉錄、65–100% 分軌。"""
    import subprocess, json, tempfile, platform, re as _re, threading
    try:
        if platform.system() == "Windows":
            os.environ["PATH"] += ";" + r"C:\Users\kamiy\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.1.1-full_build\bin"

        # ── 階段一：whisper.cpp 轉錄（0–30%，Metal 很快，給小比例）──
        binary = _resolve_whispercpp_binary()
        if not binary:
            result_q.put(("error",
                "找不到 whisper-cli。\nMac 請在終端機執行：brew install whisper-cpp\n"
                "Windows 請安裝 whisper.cpp 並將其加入系統 PATH（或放進 bin/whispercpp-win/）。"))
            return
        models_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models")
        log_q.put(f"使用引擎：whisper.cpp ＋ 聲學分軌（模型 {model_name}）")
        model_path = _ensure_whispercpp_model(model_name, models_dir, log_q)

        log_q.put(f"① 準備音檔：{os.path.basename(audio)}")
        wav = os.path.join(tempfile.gettempdir(), "voxlog_wcppdiar_input.wav")
        subprocess.run(["ffmpeg", "-hide_banner", "-loglevel", "error", "-y", "-i", audio,
                        "-ar", "16000", "-ac", "1", wav], check=True)

        base = os.path.splitext(os.path.basename(audio))[0]
        out_json_base = os.path.join(tempfile.gettempdir(), f"voxlog_wcppdiar_{base}")
        cmd = [binary, "-m", model_path, "-f", wav, "-oj", "-of", out_json_base, "-pp",
               "-l", lang or "auto"]
        if prompt:
            cmd += ["--prompt", prompt]
        log_q.put("① whisper.cpp 轉錄中…")
        proc = subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE, text=True)
        for line in proc.stderr:
            m = _re.search(r"progress\s*=\s*(\d+)%", line)
            if m:
                prog_q.put(int(int(m.group(1)) * 0.30))  # 轉錄佔 0–30%
        proc.wait()
        if proc.returncode != 0:
            result_q.put(("error", f"whisper.cpp 轉錄失敗（return code {proc.returncode}）"))
            return
        prog_q.put(30)

        with open(out_json_base + ".json", encoding="utf-8") as jf:
            raw_segs = json.load(jf).get("transcription", [])
        # 轉成 whisperX align 需要的格式（秒）
        segments = []
        for seg in raw_segs:
            off = seg.get("offsets", {})
            text = (seg.get("text") or "").strip()
            if not text:
                continue
            segments.append({"start": (off.get("from", 0) or 0) / 1000.0,
                             "end": (off.get("to", 0) or 0) / 1000.0,
                             "text": text})
        if not segments:
            result_q.put(("error", "whisper.cpp 沒有產出任何文字段落，無法分軌。"))
            return

        # ── 階段二：whisperX align + pyannote 分軌（65–100%）──
        import whisperx, torch
        if torch.cuda.is_available():
            device = "cuda"
        elif platform.system() == "Darwin" and torch.backends.mps.is_available():
            os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"  # pyannote 少數 op 沒 MPS 實作，讓它退回 CPU 而非崩潰
            device = "mps"
        else:
            device = "cpu"
        log_q.put(f"② 使用裝置：{device.upper()}（對齊／分軌）")
        audio_data = whisperx.load_audio(audio)

        log_q.put("② 對齊時間戳…")
        amodel, meta = whisperx.load_align_model(language_code=(lang or "zh"), device=device)
        aligned = whisperx.align(segments, amodel, meta, audio_data, device,
                                 return_char_alignments=False)
        prog_q.put(40)

        log_q.put("② 分析說話人（pyannote 聲紋，這段最久，請耐心）…")
        from whisperx.diarize import DiarizationPipeline
        diar = DiarizationPipeline(token=hf_token, device=device)
        dkw = {}
        if num_speakers and num_speakers > 0:
            dkw["num_speakers"] = num_speakers
        # pyannote 不吐逐步進度；用一條漸近爬升（40→97）讓進度條別卡死，完成後補滿
        _stop_creep = threading.Event()
        def _creep():
            pct = 40.0
            while not _stop_creep.wait(2.0):
                pct += (97 - pct) * 0.05
                prog_q.put(int(pct))
        threading.Thread(target=_creep, daemon=True).start()
        try:
            diar_segs = diar(audio_data, **dkw)
        finally:
            _stop_creep.set()
        result = whisperx.assign_word_speakers(diar_segs, aligned)
        prog_q.put(98)

        try:
            import opencc
            converter = opencc.OpenCC("s2twp")
        except Exception:
            converter = None

        out_path = out_path_override or os.path.join(out_dir, f"{base}_transcript.txt")
        f, out_path = _safe_open_transcript(out_path, log_q)
        with f:
            current_speaker = None
            for seg in result["segments"]:
                start = seg.get("start", 0)
                end = seg.get("end", 0)
                speaker = seg.get("speaker", "SPEAKER_??")
                text = seg.get("text", "").strip()
                if not text:
                    continue
                if converter:
                    text = converter.convert(text)
                text = _collapse_repeats(text)
                ts = f"[{int(start)//60:02d}:{int(start)%60:02d} --> {int(end)//60:02d}:{int(end)%60:02d}]"
                if speaker != current_speaker:
                    if current_speaker is not None:
                        f.write("\n")
                    f.write(f"{speaker}：\n")
                    current_speaker = speaker
                f.write(f"{ts} {text}\n")
        prog_q.put(100)

        log_q.put("分軌完成；可在第二步用「修正說話者」把 SPEAKER_xx 改成真實角色名。")
        for tmp_f in (wav, out_json_base + ".json"):
            try:
                os.remove(tmp_f)
            except OSError:
                pass
        result_q.put(("done", out_path))
    except Exception as e:
        result_q.put(("error", str(e)))


def _make_icon():
    try:
        from PIL import Image, ImageDraw
        # 4× 超取樣後再縮小 → 線條與圓角都帶抗鋸齒，不再糊
        SS = 4
        size = 256
        S = size * SS
        img = Image.new("RGBA", (S, S), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)

        # ── 圓角方形底（macOS 風格 squircle，帶垂直漸層）──
        radius = int(S * 0.22)
        bg = Image.new("RGBA", (S, S), (0, 0, 0, 0))
        bg_draw = ImageDraw.Draw(bg)
        top, bot = (44, 46, 52), (24, 24, 26)
        for y in range(S):
            t = y / S
            r = int(top[0] + (bot[0] - top[0]) * t)
            g = int(top[1] + (bot[1] - top[1]) * t)
            b = int(top[2] + (bot[2] - top[2]) * t)
            bg_draw.line([(0, y), (S, y)], fill=(r, g, b, 255))
        mask = Image.new("L", (S, S), 0)
        ImageDraw.Draw(mask).rounded_rectangle([0, 0, S - 1, S - 1], radius=radius, fill=255)
        img.paste(bg, (0, 0), mask)

        # ── 麥克風 ──
        cx     = S / 2
        accent = (74, 158, 255, 255)
        lw     = max(3, int(S * 0.046))
        half    = S * 0.115           # 麥克頭半寬
        cap_top = S * 0.20
        cap_bot = S * 0.48
        # 麥克頭（圓角膠囊）
        draw.rounded_rectangle(
            [cx - half, cap_top, cx + half, cap_bot],
            radius=half, fill=accent,
        )
        # U 形托架（弧，開口朝上讓麥克頭穿出）
        cy_cap = (cap_top + cap_bot) / 2
        r_arc  = half + S * 0.078
        draw.arc(
            [cx - r_arc, cy_cap - r_arc, cx + r_arc, cy_cap + r_arc],
            start=-20, end=200, fill=accent, width=lw,
        )
        # 支桿
        stem_top = cy_cap + r_arc
        base_y   = S * 0.77
        draw.line([(cx, stem_top), (cx, base_y)], fill=accent, width=lw)
        # 底座
        bw = S * 0.145
        draw.rounded_rectangle(
            [cx - bw, base_y - lw / 2, cx + bw, base_y + lw / 2],
            radius=lw / 2, fill=accent,
        )

        return img.resize((size, size), Image.LANCZOS)
    except Exception:
        return None


def _make_audio_icon(size=38):
    try:
        from PIL import Image, ImageDraw
        img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        cx, cy = size // 2, size // 2
        bar_w = max(2, size // 12)
        heights = [0.35, 0.6, 0.85, 1.0, 0.85, 0.6, 0.35]
        n = len(heights)
        gap = bar_w
        total = n * bar_w + (n - 1) * gap
        x0 = cx - total // 2
        color = (255, 255, 255, 210)
        for i, h in enumerate(heights):
            bh = int(size * h * 0.72)
            x = x0 + i * (bar_w + gap)
            y0_ = cy - bh // 2
            y1_ = cy + bh // 2
            draw.rounded_rectangle([x, y0_, x + bar_w, y1_], radius=bar_w // 2, fill=color)
        return img
    except Exception:
        return None


def _make_yt_icon(size=22):
    try:
        from PIL import Image, ImageDraw
        img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        pad = 1
        rh = size // 3
        draw.rounded_rectangle(
            [pad, rh, size - pad, size - rh],
            radius=size // 7, fill=(204, 0, 0, 255),
        )
        cx, cy = size // 2, size // 2
        t = size // 5
        draw.polygon([(cx - t, cy - t), (cx + t + 1, cy), (cx - t, cy + t)], fill="white")
        return img
    except Exception:
        return None


class SpeakerNameDialog(ctk.CTkToplevel):
    def __init__(self, parent, speakers, samples=None):
        super().__init__(parent)
        self.title("設定說話者名稱")
        self.resizable(False, False)
        self.configure(fg_color=BG)
        self.grab_set()
        self.result = None
        self.entries = {}
        samples = samples or {}

        ctk.CTkLabel(
            self,
            text=f"辨識出 {len(speakers)} 位說話者，請依「代表發言」判斷各是誰、填入名稱：",
            fg_color="transparent", text_color=TEXT,
            font=ctk.CTkFont(FONT_UI, 14),
        ).pack(anchor="w", padx=20, pady=(18, 4))
        ctk.CTkLabel(
            self,
            text="（不確定就先留原樣或「略過」，逐字稿裡的代表發言可幫你回頭對照）",
            fg_color="transparent", text_color=SUBTEXT,
            font=ctk.CTkFont(FONT_UI, 12),
        ).pack(anchor="w", padx=20, pady=(0, 10))

        for spk in speakers:
            row_frame = ctk.CTkFrame(self, fg_color=SURFACE, corner_radius=8)
            row_frame.pack(fill="x", padx=20, pady=5, ipady=4)
            top = ctk.CTkFrame(row_frame, fg_color="transparent")
            top.pack(fill="x", padx=10, pady=(6, 2))
            ctk.CTkLabel(top, text=spk,
                         fg_color="transparent", text_color=ACCENT,
                         font=ctk.CTkFont(FONT_UI, 13, weight="bold"),
                         width=110, anchor="w").pack(side="left", padx=(0, 10))
            var = tk.StringVar(value=spk)
            ctk.CTkEntry(top, textvariable=var, width=220,
                         placeholder_text="輸入這位的名稱…",
                         fg_color=BG, text_color=TEXT,
                         border_color=BORDER, border_width=1,
                         font=ctk.CTkFont(FONT_UI, 13)).pack(side="left")
            self.entries[spk] = var
            quote = samples.get(spk, "")
            if quote:
                ctk.CTkLabel(
                    row_frame, text=f"代表發言：「{quote}」",
                    fg_color="transparent", text_color=SUBTEXT,
                    font=ctk.CTkFont(FONT_UI, 12), justify="left",
                    wraplength=380, anchor="w",
                ).pack(fill="x", anchor="w", padx=10, pady=(0, 6))

        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=(12, 18), padx=20, fill="x")
        ctk.CTkButton(btn_frame, text="確認", command=self._ok,
                      fg_color=GREEN, hover_color="#219A52",
                      text_color="white", width=90).pack(side="right", padx=(8, 0))
        ctk.CTkButton(btn_frame, text="略過", command=self._cancel,
                      fg_color=SURFACE, hover_color=BORDER,
                      text_color=TEXT, width=90).pack(side="right")
        self.transient(parent)
        self.wait_window()

    def _ok(self):
        self.result = {spk: (var.get().strip() or spk)
                       for spk, var in self.entries.items()}
        self.destroy()

    def _cancel(self):
        self.destroy()


class PartyContextDialog(ctk.CTkToplevel):
    """按「自動標註說話方」前彈出，讓使用者補人名/情境。
    result：None=取消；字串（可能空）=確定，當作標註用的情境。"""

    def __init__(self, parent, prefill=""):
        super().__init__(parent)
        self.title("標註說話方 — 這場對話有誰？")
        self.resizable(False, False)
        self.configure(fg_color=BG)
        self.grab_set()
        self.result = None

        ctk.CTkLabel(
            self, justify="left", fg_color="transparent", text_color=TEXT,
            font=ctk.CTkFont(FONT_UI, 15, weight="bold"),
            text="告訴 AI 這場對話有誰、各代表哪一方",
        ).pack(anchor="w", padx=22, pady=(20, 2))
        ctk.CTkLabel(
            self, justify="left", fg_color="transparent", text_color=SUBTEXT,
            font=ctk.CTkFont(FONT_UI, 13), wraplength=480,
            text="想看到實際人名，就在這裡寫清楚誰是誰；標出來就會用這些名稱。\n"
                 "例如：我方＝我們公司、對方＝客戶；或：主講人＝王經理、提問者＝陳工程師。\n"
                 "留空也可以，AI 會自己依內容判斷角色（例如「主管／部屬」）。",
        ).pack(anchor="w", padx=22, pady=(0, 10))

        self.box = ctk.CTkTextbox(
            self, width=480, height=110, fg_color=SURFACE, text_color=TEXT,
            border_color=BORDER, border_width=1, font=ctk.CTkFont(FONT_UI, 13),
            wrap="word",
        )
        self.box.pack(padx=22)
        if prefill:
            self.box.insert("1.0", prefill)

        btn = ctk.CTkFrame(self, fg_color="transparent")
        btn.pack(pady=(14, 20), padx=22, fill="x")
        ctk.CTkButton(btn, text="開始標註", command=self._ok,
                      fg_color=GREEN, hover_color="#219A52", text_color="white",
                      width=110, font=ctk.CTkFont(FONT_UI, 13)).pack(side="right", padx=(8, 0))
        ctk.CTkButton(btn, text="取消", command=self._cancel,
                      fg_color=SURFACE, hover_color=BORDER, text_color=TEXT,
                      width=90, font=ctk.CTkFont(FONT_UI, 13)).pack(side="right")
        self.transient(parent)
        self.wait_window()

    def _ok(self):
        self.result = self.box.get("1.0", "end").strip()
        self.destroy()

    def _cancel(self):
        self.result = None
        self.destroy()


class YouTubeURLDialog(ctk.CTkToplevel):
    """點 YouTube 按鈕後彈出，輸入網址（Cookies 收為進階選項）。result = (url, cookies_choice) 或 None。"""

    COOKIE_VALUES = ["Cookies: 無", "Cookies: Chrome", "Cookies: Edge", "Cookies: Firefox",
                     "Cookies: Brave", "Cookies: Opera", "Cookies: Safari",
                     "Cookies: 檔案", "Cookies: 選擇檔案..."]

    def __init__(self, parent, url_default="", cookies_default="Cookies: 無"):
        super().__init__(parent)
        self.title("從 YouTube 下載音訊")
        self.resizable(False, False)
        self.configure(fg_color=BG)
        self.grab_set()
        self.result = None

        ctk.CTkLabel(self, text="貼上 YouTube 影片網址",
                     fg_color="transparent", text_color=TEXT,
                     font=ctk.CTkFont(FONT_UI, 15, weight="bold")).pack(anchor="w", padx=20, pady=(18, 6))
        self.url_var = tk.StringVar(value=url_default)
        entry = ctk.CTkEntry(self, textvariable=self.url_var, width=440, height=36,
                             placeholder_text="https://www.youtube.com/watch?v=...",
                             fg_color=SURFACE, text_color=TEXT,
                             border_color=BORDER, border_width=1,
                             font=ctk.CTkFont(FONT_UI, 13))
        entry.pack(padx=20)

        ctk.CTkLabel(self,
                     text="一般公開影片免填 Cookies；遇到年齡限制／會員／私人影片才需要。",
                     fg_color="transparent", text_color=SUBTEXT,
                     font=ctk.CTkFont(FONT_UI, 12), justify="left",
                     anchor="w").pack(anchor="w", padx=20, pady=(12, 2))
        self.cookies_var = tk.StringVar(
            value=cookies_default if cookies_default in self.COOKIE_VALUES else "Cookies: 無")
        ctk.CTkComboBox(self, variable=self.cookies_var, values=self.COOKIE_VALUES,
                        width=240, state="readonly",
                        fg_color=SURFACE, text_color=TEXT, button_color=BORDER,
                        button_hover_color=ACCENT, border_color=BORDER,
                        dropdown_fg_color=SURFACE, dropdown_text_color=TEXT,
                        dropdown_hover_color=BORDER,
                        font=ctk.CTkFont(FONT_UI, 12)).pack(anchor="w", padx=20)

        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=(16, 18), padx=20, fill="x")
        ctk.CTkButton(btn_frame, text="開始下載", command=self._ok,
                      fg_color=GREEN, hover_color="#219A52",
                      text_color="white", width=100).pack(side="right", padx=(8, 0))
        ctk.CTkButton(btn_frame, text="取消", command=self._cancel,
                      fg_color=SURFACE, hover_color=BORDER,
                      text_color=TEXT, width=90).pack(side="right")

        self.transient(parent)
        entry.focus_set()
        self.bind("<Return>", lambda *_: self._ok())
        self.bind("<Escape>", lambda *_: self._cancel())
        self.wait_window()

    def _ok(self):
        url = self.url_var.get().strip()
        if not url:
            return
        self.result = (url, self.cookies_var.get())
        self.destroy()

    def _cancel(self):
        self.destroy()


MODEL_HELP_TEXT = (
    "模型越大越準，但越慢、越吃記憶體：\n"
    "• small（約 2GB）：快；安靜環境、標準國語就夠用\n"
    "• medium（約 5GB）：較準；扛得住口音、專有名詞、中英夾雜\n"
    "• large（約 10GB）：最準但最慢\n"
    "• tiny / base：最省資源，僅測試用\n\n"
    "8GB 記憶體的電腦請用 small；錯字可在下一步\n"
    "「AI 校正逐字稿」補回。"
)


class _Tooltip:
    """滑鼠移上去顯示說明的小浮窗（tkinter 沒內建，自己做一個）。"""

    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip = None
        widget.bind("<Enter>", self._show, add="+")
        widget.bind("<Leave>", self._hide, add="+")

    def _show(self, _=None):
        if self.tip or not self.text:
            return
        x = self.widget.winfo_rootx() + self.widget.winfo_width() + 8
        y = self.widget.winfo_rooty() - 2
        self.tip = tk.Toplevel(self.widget)
        self.tip.wm_overrideredirect(True)
        self.tip.wm_geometry(f"+{x}+{y}")
        try:
            self.tip.attributes("-topmost", True)
        except Exception:
            pass
        # 外框用 BORDER、內層 SURFACE，做出細邊框效果
        outer = tk.Frame(self.tip, bg=BORDER)
        outer.pack()
        tk.Label(outer, text=self.text, justify="left",
                 bg=SURFACE, fg=TEXT, font=(FONT_UI, 15),
                 padx=14, pady=12).pack(padx=1, pady=1)

    def _hide(self, _=None):
        if self.tip:
            self.tip.destroy()
            self.tip = None


class TranscribeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("VoxLog")
        self.root.geometry("780x720")
        self.root.minsize(720, 660)
        self.root.configure(fg_color=BG)

        self.root.grid_columnconfigure(0, weight=1)
        # 分頁區「貼合內容」（weight 0 → 取自身需求高度，開始鈕永不被切）；
        # 多出來的高度給 Log（row 5 weight 1），視窗拉高就是 Log 變清楚，不再留空隙。
        self.root.grid_rowconfigure(0, weight=0)
        self.root.grid_rowconfigure(5, weight=1)

        self.process = None
        self.result_q = None
        self.log_q = None
        self.prog_q = None
        self.timer_start = None
        self.timer_id = None
        self.transcript_path = None
        self.notes_thread = None
        self.ai_cancelled = False
        self._ai_anim_running = False
        self._ai_anim_val = 0.0
        self._ai_anim_dir = 1
        self.audio_var = tk.StringVar()
        self.out_var = tk.StringVar()
        self.srt_var = tk.BooleanVar(value=True)
        ensure_config()  # 首次啟動自動建立空白 config.json，小白不必手動建檔
        self.cfg = load_config()
        self.cookies_file_path = self.cfg.get("cookies_file_path", "")
        # YouTube 下載相關（改用對話框輸入，故在此建立變數，不在版面上佔位）
        self.yt_url_var = tk.StringVar()
        self.cookies_browser_var = tk.StringVar(value=self.cfg.get("cookies_browser", "Cookies: 無"))

        self._set_icon()
        self._build_ui()
        self.root.after(120, self._bring_to_front)

    def _bring_to_front(self):
        """啟動時把視窗帶到前景並取得焦點（macOS 上 python 跑的視窗預設不會搶前景）。"""
        try:
            self.root.lift()
            self.root.attributes("-topmost", True)
            self.root.after(300, lambda: self.root.attributes("-topmost", False))
            self.root.focus_force()
        except Exception:
            pass
        if platform.system() == "Darwin":
            try:
                from AppKit import NSApplication
                NSApplication.sharedApplication().activateIgnoringOtherApps_(True)
            except Exception:
                pass

    def _set_icon(self):
        img = _make_icon()
        if img is None:
            return
        try:
            from PIL import ImageTk, Image
            self._icon_photos = [
                ImageTk.PhotoImage(img.resize((s, s), Image.LANCZOS))
                for s in [16, 32, 64, 128, 256]
            ]
            self.root.iconphoto(True, *self._icon_photos)
        except Exception:
            pass

        # macOS：iconphoto 吃不到 Dock，改用 NSApplication 直接設定 Dock 圖示
        if platform.system() == "Darwin":
            try:
                import io
                from AppKit import NSApplication, NSImage
                from Foundation import NSData
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                data = NSData.dataWithBytes_length_(buf.getvalue(), len(buf.getvalue()))
                ns_img = NSImage.alloc().initWithData_(data)
                NSApplication.sharedApplication().setApplicationIconImage_(ns_img)
            except Exception:
                pass

    # ── Section divider ────────────────────────────
    def _section_divider(self, row, text, pady=(14, 6)):
        frame = ctk.CTkFrame(self.root, fg_color="transparent")
        frame.grid(row=row, column=0, sticky="we", padx=20, pady=pady)
        frame.grid_columnconfigure(0, weight=1)
        frame.grid_columnconfigure(2, weight=1)
        ctk.CTkFrame(frame, height=1, fg_color=BORDER).grid(
            row=0, column=0, sticky="we", padx=(0, 10), pady=7)
        ctk.CTkLabel(frame, text=text,
                     fg_color="transparent", text_color=SUBTEXT,
                     font=ctk.CTkFont(FONT_UI, 11)).grid(row=0, column=1)
        ctk.CTkFrame(frame, height=1, fg_color=BORDER).grid(
            row=0, column=2, sticky="we", padx=(10, 0), pady=7)

    # ── Build UI ───────────────────────────────────
    def _build_ui(self):
        F = ctk.CTkFont

        # ── 分頁：① 轉錄 / ② 逐字稿處理（加圖示讓分頁更醒目）──
        self._TAB_TRANSCRIBE = "🎙️  ①  轉錄逐字稿"
        self._TAB_AI = "📝  ②  逐字稿處理（AI 校正 / 摘要 / 字幕檔）"
        self.tabs = ctk.CTkTabview(
            self.root, fg_color=BG,
            segmented_button_fg_color=SURFACE,
            segmented_button_selected_color=ACCENT,
            segmented_button_selected_hover_color=ACCENT,
            segmented_button_unselected_color=SURFACE,
            segmented_button_unselected_hover_color=BORDER,
            text_color=TEXT,
        )
        self.tabs.grid(row=0, column=0, sticky="nswe", padx=16, pady=(10, 2))
        # 放大分頁按鈕，讓使用者清楚看到目前在哪一頁
        try:
            self.tabs._segmented_button.configure(font=F(FONT_UI, 16, weight="bold"), height=40)
        except Exception:
            pass
        tab_t = self.tabs.add(self._TAB_TRANSCRIBE)
        tab_ai = self.tabs.add(self._TAB_AI)

        # ════════ 分頁一：轉錄 ════════
        ctk.CTkLabel(
            tab_t, justify="left",
            text="① 選音檔或貼 YouTube 網址　② 設好輸出資料夾　③ 按「開始轉錄」",
            fg_color="transparent", text_color=TEXT, font=F(FONT_UI, 16),
        ).pack(anchor="w", padx=4, pady=(6, 4))

        # ── 辨識設定（引擎 / 語言 + 隨引擎變動的選項）──
        settings = ctk.CTkFrame(tab_t, fg_color=SURFACE, corner_radius=8)
        settings.pack(fill="x", pady=(0, 4), ipady=6)

        # 第一列：辨識引擎 + 語言
        row_eng = ctk.CTkFrame(settings, fg_color="transparent")
        row_eng.pack(fill="x", padx=14, pady=(6, 2))
        ctk.CTkLabel(row_eng, text="辨識引擎", fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 14)).pack(side="left", padx=(0, 6))
        # 沒存過引擎（首次啟動）才自動偵測：MacBook Air / 8GB → 自動選 whisper.cpp
        _saved_engine = self.cfg.get("stt_engine")
        if not _saved_engine:
            _saved_engine = _suggest_stt_engine()
            save_config({"stt_engine": _saved_engine})
            self.cfg["stt_engine"] = _saved_engine
            # whisper.cpp 在 8GB 上跑 medium 也只吃 ~1.3GB 卻準很多，故預設 medium
            # （whisper/whisperx 不在此設，維持依記憶體建議的 small）
            if _saved_engine == "whispercpp" and not self.cfg.get("stt_model"):
                save_config({"stt_model": "medium"})
                self.cfg["stt_model"] = "medium"
        self.engine_var = tk.StringVar(
            value=_STT_VAL2LABEL.get(_saved_engine, _STT_VAL2LABEL["whisper"]))
        ctk.CTkComboBox(row_eng, variable=self.engine_var,
                        values=[lb for _, lb in STT_ENGINES],
                        width=250, state="readonly",
                        command=lambda *_: self._on_engine_change(),
                        fg_color=BG, text_color=TEXT, button_color=BORDER,
                        button_hover_color=ACCENT, border_color=BORDER,
                        dropdown_fg_color=SURFACE, dropdown_text_color=TEXT,
                        dropdown_hover_color=BORDER,
                        font=F(FONT_UI, 13)).pack(side="left")

        ctk.CTkLabel(row_eng, text="語言", fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 14)).pack(side="left", padx=(18, 4))
        self.lang_var = tk.StringVar(value="中文")
        ctk.CTkComboBox(row_eng, variable=self.lang_var,
                        values=list(LANG_MAP.keys()),
                        width=110, state="readonly",
                        fg_color=BG, text_color=TEXT, button_color=BORDER,
                        button_hover_color=ACCENT, border_color=BORDER,
                        dropdown_fg_color=SURFACE, dropdown_text_color=TEXT,
                        dropdown_hover_color=BORDER,
                        font=F(FONT_UI, 13)).pack(side="left")

        # 第二列：隨引擎變動的選項（地端模型 / WhisperX 額外的 HF Token + 說話人數）
        self.stt_opt_row = ctk.CTkFrame(settings, fg_color="transparent")
        self.stt_opt_row.pack(fill="x", padx=14, pady=(2, 4))

        # ▸ 地端：轉錄模型
        self.f_model = ctk.CTkFrame(self.stt_opt_row, fg_color="transparent")
        ctk.CTkLabel(self.f_model, text="轉錄模型", fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 14)).pack(side="left", padx=(0, 4))
        self.model_var = tk.StringVar(value=self.cfg.get("stt_model") or _suggest_model())
        ctk.CTkComboBox(self.f_model, variable=self.model_var,
                        values=["tiny", "base", "small", "medium", "large", "large-v3"],
                        width=120, state="readonly",
                        fg_color=BG, text_color=TEXT, button_color=BORDER,
                        button_hover_color=ACCENT, border_color=BORDER,
                        dropdown_fg_color=SURFACE, dropdown_text_color=TEXT,
                        dropdown_hover_color=BORDER,
                        font=F(FONT_UI, 13)).pack(side="left")
        _model_info = ctk.CTkLabel(self.f_model, text="ⓘ", fg_color="transparent",
                                   text_color=ACCENT, font=F(FONT_UI, 15), cursor="hand2")
        _model_info.pack(side="left", padx=(4, 0))
        _Tooltip(_model_info, MODEL_HELP_TEXT)
        self.model_var.trace_add("write", self._on_stt_model_change)

        # ▸ WhisperX 額外：HF Token + 說話人數
        self.f_wx = ctk.CTkFrame(self.stt_opt_row, fg_color="transparent")
        ctk.CTkLabel(self.f_wx, text="HF Token", fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 14)).pack(side="left", padx=(16, 6))
        self.token_var = tk.StringVar(value=self.cfg.get("hf_token", ""))
        ctk.CTkEntry(self.f_wx, textvariable=self.token_var, width=210, show="*",
                     fg_color=BG, text_color=TEXT, border_color=BORDER, border_width=1,
                     font=F(FONT_UI, 13)).pack(side="left", padx=(0, 12))
        ctk.CTkLabel(self.f_wx, text="說話人數", fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 14)).pack(side="left", padx=(0, 6))
        self.speakers_var = tk.IntVar(value=0)
        tk.Spinbox(self.f_wx, from_=0, to=10, textvariable=self.speakers_var,
                   width=4, bg=SURFACE, fg=TEXT, buttonbackground=BORDER,
                   relief="flat", font=(FONT_UI, 13),
                   highlightthickness=1, highlightbackground=BORDER).pack(side="left")
        ctk.CTkLabel(self.f_wx, text="（0 ＝ 自動）", fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 13)).pack(side="left", padx=(6, 0))

        # 來源選擇（兩個方形按鈕，置中）
        source = ctk.CTkFrame(tab_t, fg_color="transparent")
        source.pack(pady=(6, 3))

        audio_col = ctk.CTkFrame(source, fg_color="transparent")
        audio_col.pack(side="left", padx=24, anchor="n")
        audio_img = _make_audio_icon(28)
        self._audio_icon = ctk.CTkImage(
            light_image=audio_img, dark_image=audio_img, size=(28, 28)
        ) if audio_img else None
        ctk.CTkButton(
            audio_col, text="選擇音/影檔", image=self._audio_icon, compound="top",
            command=self.pick_audio,
            fg_color=ACCENT, hover_color="#2980B9", text_color="white",
            font=F(FONT_UI, 14, weight="bold"), width=150, height=48,
        ).pack()
        self.audio_display_var = tk.StringVar(value="未選擇")
        self.audio_display_label = ctk.CTkLabel(
            audio_col, textvariable=self.audio_display_var,
            fg_color="transparent", text_color=SUBTEXT,
            font=F(FONT_UI, 15, weight="bold"), width=200, wraplength=220,
            justify="center", anchor="center")
        self.audio_display_label.pack(pady=(6, 0))

        yt_col = ctk.CTkFrame(source, fg_color="transparent")
        yt_col.pack(side="left", padx=24, anchor="n")
        yt_img = _make_yt_icon(28)
        self._yt_icon = ctk.CTkImage(
            light_image=yt_img, dark_image=yt_img, size=(28, 28)
        ) if yt_img else None
        self.yt_btn = ctk.CTkButton(
            yt_col, text="YouTube下載", image=self._yt_icon, compound="top",
            command=self.open_youtube_dialog,
            fg_color="#CC0000", hover_color="#990000", text_color="white",
            font=F(FONT_UI, 14, weight="bold"), width=150, height=48,
        )
        self.yt_btn.pack()
        ctk.CTkLabel(yt_col, text="點一下貼上網址",
                     fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 12), width=150, anchor="center").pack(pady=(4, 0))

        # 切到正確的引擎，顯示對應的設定列
        self._on_engine_change()

        # ── 輸出資料夾 + 字幕（移到情境上方，維持「設輸出 → 填情境 → 開始」順序）──
        r_out = ctk.CTkFrame(tab_t, fg_color="transparent")
        self.r_out = r_out
        r_out.pack(fill="x", pady=(6, 4))
        ctk.CTkButton(r_out, text="輸出資料夾", command=self.pick_output,
                      fg_color=SURFACE, hover_color=BORDER, text_color=TEXT,
                      font=F(FONT_UI, 14), width=110).pack(side="left")
        self.out_display_var = tk.StringVar(value="未選擇")
        ctk.CTkLabel(r_out, textvariable=self.out_display_var,
                     fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 13), anchor="w",
                     width=180).pack(side="left", padx=(10, 0))
        ctk.CTkCheckBox(r_out, text="輸出字幕 SRT", variable=self.srt_var,
                        fg_color=ACCENT, hover_color="#2980B9",
                        text_color=TEXT, font=F(FONT_UI, 13),
                        ).pack(side="left", padx=(16, 0))

        # ── 情境（左 65%）＋ 開始/停止（右側、上下排列）──
        ctx_row = ctk.CTkFrame(tab_t, fg_color="transparent")
        ctx_row.pack(fill="x", pady=(4, 4))
        ctx_row.grid_columnconfigure(0, weight=1)   # 情境欄吃滿剩餘寬度
        ctx_row.grid_columnconfigure(1, weight=0)   # 按鈕欄＝固定窄寬（內容自身寬度）

        ctx_t = ctk.CTkFrame(ctx_row, fg_color="transparent")
        ctx_t.grid(row=0, column=0, sticky="nsew", padx=(0, 14))
        term_head = ctk.CTkFrame(ctx_t, fg_color="transparent")
        term_head.pack(fill="x")
        ctk.CTkLabel(term_head, text="關鍵術語 / 人名（用逗號分隔）", fg_color="transparent",
                     text_color=TEXT, font=F(FONT_UI, 14)).pack(side="left")
        self.term_budget_var = tk.StringVar(value="")
        self.term_budget_label = ctk.CTkLabel(
            term_head, textvariable=self.term_budget_var, fg_color="transparent",
            text_color=SUBTEXT, font=F(FONT_UI, 12))
        self.term_budget_label.pack(side="left", padx=(8, 0))
        # 每場會議不同，故不預帶上次輸入（不持久化）
        self.terms_var = tk.StringVar(value="")
        self.terms_entry = ctk.CTkEntry(
            ctx_t, textvariable=self.terms_var,
            placeholder_text="例：鴻海, 俊傑, 依芬, CDP, 案號 A123（客戶/與會者/產品/案號，提示模型認對專有名詞）",
            fg_color=SURFACE, text_color=TEXT, border_color=BORDER, border_width=1,
            font=F(FONT_UI, 13))
        self.terms_entry.pack(fill="x", pady=(2, 6))
        self.terms_var.trace_add("write", self._update_term_budget)

        ctk.CTkLabel(ctx_t, text="會議背景（給 AI 校正/摘要參考，不限字數）",
                     fg_color="transparent", text_color=TEXT,
                     font=F(FONT_UI, 14)).pack(anchor="w")
        self.bg_box = ctk.CTkTextbox(ctx_t, height=48, fg_color=SURFACE, text_color=TEXT,
                                     border_color=BORDER, border_width=1, font=F(FONT_UI, 13),
                                     wrap="word")
        self.bg_box.pack(fill="x", pady=(2, 0))

        # 右側：開始 / 停止（窄固定寬、上下排列、撐滿情境區高度，好按）
        btn_col = ctk.CTkFrame(ctx_row, fg_color="transparent")
        btn_col.grid(row=0, column=1, sticky="nsew")
        self.start_btn = ctk.CTkButton(
            btn_col, text="▶ 開始轉錄", command=self.start, width=108,
            fg_color=GREEN, hover_color="#219A52", text_color="white",
            font=F(FONT_UI, 14, weight="bold"))
        self.start_btn.pack(fill="both", expand=True, pady=(18, 4))
        self.stop_btn = ctk.CTkButton(
            btn_col, text="■ 停止", command=self.stop, width=108,
            fg_color=RED, hover_color=RED_HOVER, text_color="white",
            font=F(FONT_UI, 14, weight="bold"), state="disabled")
        self.stop_btn.pack(fill="both", expand=True, pady=(4, 0))

        self._update_term_budget()

        # ════════ 分頁二：AI 後製 ════════
        ctk.CTkLabel(
            tab_ai, justify="left",
            text="轉錄完成會自動帶你到這頁；也可直接按下方「載入逐字稿」處理舊檔。\n"
                 "① 選 AI 引擎　② 填 API Key 按「驗證」變綠　③ 點「校正逐字稿 / 產生摘要 / 匯出 SRT」",
            fg_color="transparent", text_color=TEXT, font=F(FONT_UI, 16),
        ).pack(anchor="w", padx=4, pady=(14, 12))

        ai_frame = ctk.CTkFrame(tab_ai, fg_color="transparent")
        ai_frame.pack(fill="x", pady=(0, 4))

        # 第一行：AI 引擎選擇
        ai_engine_row = ctk.CTkFrame(ai_frame, fg_color="transparent")
        ai_engine_row.pack(fill="x", anchor="w")
        ctk.CTkLabel(ai_engine_row, text="AI 引擎",
                     fg_color="transparent", text_color=TEXT,
                     font=F(FONT_UI, 14), width=72, anchor="e").pack(side="left", padx=(0, 8))
        self.ai_engine_var = tk.StringVar(value="gemini")
        for val, label in [("claude", "Claude"), ("gemini", "Gemini"), ("openai", "ChatGPT"), ("ollama", "Ollama"), ("lmstudio", "LM Studio")]:
            ctk.CTkRadioButton(
                ai_engine_row, text=label, variable=self.ai_engine_var,
                value=val, command=self._on_ai_engine_change,
                text_color=TEXT, fg_color=ACCENT, hover_color=ACCENT,
                font=F(FONT_UI, 14),
            ).pack(side="left", padx=(0, 12))

        # 第二行：模型選擇（雲端引擎才顯示；可下拉點選，也可手動輸入）
        self.ai_model_row = ctk.CTkFrame(ai_frame, fg_color="transparent")
        ctk.CTkLabel(self.ai_model_row, text="AI 模型",
                     fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 13), width=72, anchor="e").pack(side="left", padx=(0, 8))
        self.ai_model_var = tk.StringVar(
            value=self.cfg.get("gemini_model", _ENGINE_MODEL_DEFAULT["gemini"]))
        self.ai_model_combo = ctk.CTkComboBox(
            self.ai_model_row, variable=self.ai_model_var, values=_ENGINE_MODELS["gemini"],
            width=380,
            fg_color=SURFACE, text_color=TEXT, border_color=BORDER, border_width=1,
            button_color=BORDER, button_hover_color=ACCENT,
            dropdown_fg_color=SURFACE, dropdown_text_color=TEXT, dropdown_hover_color=BORDER,
            font=F(FONT_UI, 13),
        )
        self.ai_model_combo.pack(side="left", padx=(0, 4))
        self.ai_model_row.pack(fill="x", anchor="w", pady=(6, 0))

        # 第三行：API Key 輸入
        self.ai_key_row = ctk.CTkFrame(ai_frame, fg_color="transparent")
        self.ai_key_row.pack(fill="x", anchor="w", pady=(6, 0))
        self.ai_key_label = ctk.CTkLabel(self.ai_key_row, text="Gemini Key",
                                          fg_color="transparent", text_color=SUBTEXT,
                                          font=F(FONT_UI, 13), width=72, anchor="e")
        self.ai_key_label.pack(side="left", padx=(0, 8))
        self.ai_key_var = tk.StringVar(value=self.cfg.get("gemini_key", ""))
        self.verify_btn = ctk.CTkButton(
            self.ai_key_row, text="驗證", command=self._verify_api,
            font=F(FONT_UI, 13, weight="bold"), width=64,
        )
        self.verify_btn.pack(side="right", padx=(8, 0))
        self.ai_key_entry = ctk.CTkEntry(self.ai_key_row, textvariable=self.ai_key_var, show="*",
                                         fg_color=SURFACE, text_color=TEXT,
                                         border_color=BORDER, border_width=1,
                                         font=F(FONT_UI, 13))
        self.ai_key_entry.pack(side="left", fill="x", expand=True, padx=(0, 4))
        # 有填入值時才把「驗證」鈕從暗色轉成藍色，明確提示可按
        self.ai_key_var.trace_add("write", self._update_verify_btn)
        self._update_verify_btn()
        # 模型下拉變動（點選或手打）即存進 config，並依目前引擎同步顯示
        self.ai_model_var.trace_add("write", self._on_model_change)
        self._sync_model_widget(self.ai_engine_var.get())

        # 第一步：載入逐字稿（左側大按鈕，使用者一眼知道從哪開始）＋ 描述情境（右側，窄而高）
        top_row = ctk.CTkFrame(tab_ai, fg_color="transparent")
        top_row.pack(fill="x", pady=(12, 0))
        top_row.grid_columnconfigure(0, weight=0)
        top_row.grid_columnconfigure(1, weight=1)
        self.load_btn = ctk.CTkButton(
            top_row, text="載入逐字稿", command=self.load_transcript,
            fg_color=ACCENT, hover_color="#2980B9", text_color="white",
            font=F(FONT_UI, 16, weight="bold"), width=160,
        )
        self.load_btn.grid(row=0, column=0, sticky="nsew", padx=(0, 14))

        ctx_col = ctk.CTkFrame(top_row, fg_color="transparent")
        ctx_col.grid(row=0, column=1, sticky="nsew")
        ctk.CTkLabel(
            ctx_col,
            text="描述這段錄音的情境，幫助 AI 校正名稱、辨識角色；例如：鴻海的 ERP 專案會議，與會者有工程師 Joe、Rachel 與專案經理 Sam。",
            fg_color="transparent", text_color=SUBTEXT, font=F(FONT_UI, 13),
            justify="left", anchor="w", wraplength=540,
        ).pack(fill="x", anchor="w", pady=(0, 4))
        self.ai_context_box = ctk.CTkTextbox(
            ctx_col, height=92, fg_color=SURFACE, text_color=TEXT,
            border_color=BORDER, border_width=1, font=F(FONT_UI, 13),
            wrap="word",
        )
        self.ai_context_box.pack(fill="x")

        # 後製按鈕
        post_outer = ctk.CTkFrame(tab_ai, fg_color="transparent")
        post_outer.pack(fill="x", pady=(10, 4))

        # 目前處理的檔案：第一步驟完成會自動帶過來，這個標示一定要大且清楚
        self.loaded_file_var = tk.StringVar(value="尚未載入逐字稿")
        self.loaded_file_label = ctk.CTkLabel(
            post_outer, textvariable=self.loaded_file_var,
            fg_color=SURFACE, corner_radius=8,
            text_color=SUBTEXT, font=F(FONT_UI, 16, weight="bold"),
        )
        self.loaded_file_label.pack(fill="x", padx=20, pady=(10, 10), ipady=8)

        # 第二排：四個處理動作（校正 → 標出說話方 → 摘要 → SRT）
        post_row = ctk.CTkFrame(post_outer, fg_color="transparent")
        post_row.pack()
        self.correct_btn = ctk.CTkButton(
            post_row, text="校正逐字稿", command=self.correct_transcript,
            fg_color=BLUE_DIM, hover_color=BLUE, text_color=SUBTEXT,
            font=F(FONT_UI, 14), width=130, state="disabled",
        )
        self.correct_btn.pack(side="left", padx=8)
        self.rename_spk_btn = ctk.CTkButton(
            post_row, text="修正說話者", command=self.rename_speakers,
            fg_color=BLUE_DIM, hover_color=BLUE, text_color=SUBTEXT,
            font=F(FONT_UI, 14), width=120, state="disabled",
        )
        self.rename_spk_btn.pack(side="left", padx=8)
        self.notes_btn = ctk.CTkButton(
            post_row, text="產生摘要", command=self.generate_notes,
            fg_color=BLUE_DIM, hover_color=BLUE, text_color=SUBTEXT,
            font=F(FONT_UI, 14), width=120, state="disabled",
        )
        self.notes_btn.pack(side="left", padx=8)
        self.export_srt_btn = ctk.CTkButton(
            post_row, text="匯出 SRT", command=self.export_srt,
            fg_color=BLUE_DIM, hover_color=BLUE, text_color=SUBTEXT,
            font=F(FONT_UI, 14), width=110, state="disabled",
        )
        self.export_srt_btn.pack(side="left", padx=8)

        # ════════ 共用底部：計時 / 狀態 / 進度 / Log（兩個分頁都看得到）════════
        # 計時器 + 百分比
        info_frame = ctk.CTkFrame(self.root, fg_color="transparent")
        info_frame.grid(row=1, column=0, pady=(2, 2))
        self.timer_var = tk.StringVar(value="00:00:00")
        ctk.CTkLabel(info_frame, textvariable=self.timer_var,
                     fg_color="transparent", text_color=ACCENT,
                     font=F("Consolas", 18, weight="bold")).pack(side="left", padx=(0, 16))
        self.pct_var = tk.StringVar(value="")
        ctk.CTkLabel(info_frame, textvariable=self.pct_var,
                     fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 16)).pack(side="left")

        # 狀態
        self.status_label = ctk.CTkLabel(
            self.root, text="等待中...",
            fg_color="transparent", text_color=SUBTEXT,
            font=F(FONT_UI, 15, weight="bold"),
        )
        self.status_label.grid(row=2, column=0)

        # 進度條
        self.progress = ctk.CTkProgressBar(
            self.root, mode="determinate",
            fg_color=SURFACE, progress_color=ACCENT,
        )
        self.progress.set(0)
        self.progress.grid(row=3, column=0, padx=20, pady=3, sticky="we")

        # Log header
        log_hdr = ctk.CTkFrame(self.root, fg_color="transparent")
        log_hdr.grid(row=4, column=0, sticky="we", padx=20, pady=(4, 0))
        log_hdr.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(log_hdr, text="Log",
                     fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 12)).grid(row=0, column=0, sticky="w")
        ctk.CTkLabel(log_hdr, text="Kamiyu Lab",
                     fg_color="transparent", text_color=SUBTEXT,
                     font=F(FONT_UI, 13)).grid(row=0, column=2, sticky="e")

        # Log（固定小尺寸，不再吃掉版面）
        self.log = ctk.CTkTextbox(
            self.root,
            fg_color=SURFACE, text_color=TEXT,
            font=F("Consolas", 13),
            height=90,
            state="disabled",
        )
        self.log.grid(row=5, column=0, padx=20, pady=(2, 12), sticky="nswe")

    # ── YouTube 下載 ───────────────────────────────
    def open_youtube_dialog(self):
        """點 YouTube 按鈕 → 彈出對話框輸入網址（與 Cookies），確認後直接下載。"""
        dlg = YouTubeURLDialog(self.root, self.yt_url_var.get(), self.cookies_browser_var.get())
        if not dlg.result:
            return
        url, cookies_choice = dlg.result
        # 需要選 Cookies 檔案的兩種情況：先讓使用者挑檔
        if cookies_choice == "Cookies: 選擇檔案...":
            self.pick_cookies_file()
            cookies_choice = self.cookies_browser_var.get()
        elif cookies_choice == "Cookies: 檔案" and (
                not self.cookies_file_path or not os.path.exists(self.cookies_file_path)):
            self.pick_cookies_file()
            cookies_choice = self.cookies_browser_var.get()
        self.yt_url_var.set(url)
        self.cookies_browser_var.set(cookies_choice)
        self.download_youtube()

    def download_youtube(self):
        url = self.yt_url_var.get().strip()
        if not url:
            self._set_status("請輸入 YouTube 網址", RED)
            return
        out_dir = self.out_var.get().strip()
        if not out_dir:
            out_dir = os.path.join(os.path.expanduser("~"), "Downloads")
            self.out_var.set(out_dir)
            self.out_display_var.set(out_dir)

        browser_selected = self.cookies_browser_var.get()
        save_config({"cookies_browser": browser_selected})
        self.cfg["cookies_browser"] = browser_selected

        self.yt_btn.configure(state="disabled")
        self._set_status("下載中...", ACCENT)
        self.log_write(f"開始下載：{url}")
        if browser_selected == "Cookies: 檔案":
            self.log_write(f"使用 Cookies 檔案：{os.path.basename(self.cookies_file_path)}")
        elif browser_selected != "Cookies: 無":
            self.log_write(f"使用瀏覽器餅乾：{browser_selected}")

        def run():
            try:
                import yt_dlp
                final_path = [None]

                def hook(d):
                    if d["status"] == "finished":
                        base = os.path.splitext(d["filename"])[0]
                        final_path[0] = base + ".mp3"

                ydl_opts = {
                    "format": "bestaudio/best",
                    "outtmpl": os.path.join(out_dir, "%(title)s.%(ext)s"),
                    "postprocessors": [{"key": "FFmpegExtractAudio", "preferredcodec": "mp3"}],
                    "progress_hooks": [hook],
                    "quiet": True,
                    # YouTube (2026)：預設網頁客戶端需要解 JS「n 簽章」挑戰，
                    # 否則會出現 "Requested format is not available"。改用 android_vr
                    # 客戶端可免挑戰、免 cookies 直接取得音訊格式，最穩定。
                    "extractor_args": {"youtube": {"player_client": ["android_vr"]}},
                }
                # 決定要不要帶 cookies（公開影片通常不需要）
                cookie_opt = None
                if browser_selected == "Cookies: 檔案":
                    if self.cookies_file_path and os.path.exists(self.cookies_file_path):
                        cookie_opt = ("cookiefile", self.cookies_file_path)
                else:
                    browser_map = {
                        "Cookies: Chrome": "chrome",
                        "Cookies: Edge": "edge",
                        "Cookies: Firefox": "firefox",
                        "Cookies: Brave": "brave",
                        "Cookies: Opera": "opera",
                        "Cookies: Safari": "safari",
                    }
                    browser_name = browser_map.get(browser_selected)
                    if browser_name:
                        cookie_opt = ("cookiesfrombrowser", (browser_name,))

                def _do_download(opts):
                    with yt_dlp.YoutubeDL(opts) as ydl:
                        ydl.download([url])

                try:
                    opts = dict(ydl_opts)
                    if cookie_opt:
                        opts[cookie_opt[0]] = cookie_opt[1]
                    _do_download(opts)
                except Exception as e:
                    msg = str(e).lower()
                    # cookies 讀取失敗（如該瀏覽器未安裝 / 找不到設定檔）→ 自動改用無 cookies 重試
                    cookie_failed = cookie_opt and any(
                        k in msg for k in ("cookie", "could not find", "database", "profile")
                    )
                    if cookie_failed:
                        self.root.after(0, lambda: self.log_write(
                            "Cookies 讀取失敗（該瀏覽器可能未安裝），改用無 cookies 重新下載…"))
                        _do_download(dict(ydl_opts))
                    else:
                        raise
                self.root.after(0, lambda: self._yt_done(final_path[0]))
            except ImportError:
                self.root.after(0, lambda: self._yt_error("請先安裝 yt-dlp：pip install yt-dlp"))
            except Exception as e:
                err = str(e)
                self.root.after(0, lambda: self._yt_error(err))

        threading.Thread(target=run, daemon=True).start()

    def _yt_done(self, path):
        self.yt_btn.configure(state="normal")
        if path and os.path.exists(path):
            self.audio_var.set(path)
            self.audio_display_var.set(os.path.basename(path))
            self.audio_display_label.configure(text_color=ACCENT)
            self.log_write(f"下載完成：{os.path.basename(path)}")
            self._set_status("下載完成，可開始轉錄", GREEN)
        else:
            self.log_write("下載完成，請手動選擇音檔")
            self._set_status("下載完成", ACCENT)

    def _yt_error(self, err):
        self.yt_btn.configure(state="normal")
        self.log_write(f"下載錯誤：{err}")
        self._set_status("下載失敗", RED)

    # ── Callbacks ──────────────────────────────────
    def _engine_val(self):
        """目前選的辨識引擎內部值（whisper / whisperx）。"""
        return _STT_LABEL2VAL.get(self.engine_var.get(), "whisper")

    def _on_engine_change(self):
        val = self._engine_val()
        for fr in (self.f_model, self.f_wx):
            fr.pack_forget()
        self.f_model.pack(side="left")
        if val in ("whisperx", "whispercpp_diar"):
            self.f_wx.pack(side="left")
        save_config({"stt_engine": val})
        self.cfg["stt_engine"] = val

    def _on_stt_model_change(self, *_):
        """記住使用者選的地端模型，下次啟動沿用。"""
        if not hasattr(self, "model_var"):
            return
        m = self.model_var.get()
        if m and self.cfg.get("stt_model") != m:
            save_config({"stt_model": m})
            self.cfg["stt_model"] = m

    def _update_term_budget(self, *_):
        """即時顯示「關鍵術語」還剩多少 token 預算；超出就紅字提醒會被裁切。"""
        if not hasattr(self, "term_budget_var"):
            return
        terms = self.terms_var.get().strip()
        left = _TERMS_BUDGET - len(terms)
        if left >= 0:
            self.term_budget_var.set(f"（給辨識・預算還剩 {left} 字）")
            self.term_budget_label.configure(text_color=SUBTEXT)
        else:
            self.term_budget_var.set(f"（超出 {-left} 字，超出部分不會送進地端辨識）")
            self.term_budget_label.configure(text_color=RED)

    def _sync_model_widget(self, engine):
        """切換引擎時更新模型下拉：雲端引擎顯示對應清單與已存值；
        本地引擎（ollama / lmstudio）的模型即 Key 欄位，故隱藏下拉。"""
        if not hasattr(self, "ai_model_combo"):
            return
        if engine in _ENGINE_MODELS:
            labels = _ENGINE_MODELS[engine]
            self.ai_model_combo.configure(values=labels)
            cfg_key = _ENGINE_MODEL_KEY[engine]
            saved = self.cfg.get(cfg_key, _ENGINE_MODEL_DEFAULT[engine])
            # 存的是純模型 ID，顯示時換回含「免費/付費」標註的對應標籤
            display = next((lb for lb in labels if _model_id(lb) == saved), saved)
            self.ai_model_var.set(display)
            if not self.ai_model_row.winfo_ismapped():
                self.ai_model_row.pack(fill="x", anchor="w", pady=(6, 0),
                                       before=self.ai_key_row)
        else:
            self.ai_model_row.pack_forget()

    def _on_model_change(self, *_):
        """模型下拉變動（點選或手動輸入）即存進 config 對應欄位。"""
        if not hasattr(self, "ai_model_combo"):
            return
        engine = self.ai_engine_var.get()
        if engine not in _ENGINE_MODEL_KEY:
            return
        model = _model_id(self.ai_model_var.get())  # 去掉「免費/付費」標註，只存純 ID
        if not model:
            return
        cfg_key = _ENGINE_MODEL_KEY[engine]
        if self.cfg.get(cfg_key) == model:  # 沒變就不重複寫檔（含程式設定值時）
            return
        save_config({cfg_key: model})
        self.cfg[cfg_key] = model

    def _on_ai_engine_change(self):
        engine = self.ai_engine_var.get()
        if hasattr(self, "verify_btn"):  # 切換引擎時清掉上次的驗證結果
            self.verify_btn.configure(text="驗證")
        self._sync_model_widget(engine)
        if engine == "claude":
            self.ai_key_label.configure(text="Claude Key")
            self.ai_key_var.set(self.cfg.get("anthropic_key", ""))
            self.ai_key_entry.configure(show="*")
        elif engine == "openai":
            self.ai_key_label.configure(text="OpenAI Key")
            self.ai_key_var.set(self.cfg.get("openai_key", ""))
            self.ai_key_entry.configure(show="*")
        elif engine == "ollama":
            self.ai_key_label.configure(text="Model 名稱")
            self.ai_key_var.set(self.cfg.get("ollama_model", "qwen2.5:7b"))
            self.ai_key_entry.configure(show="")
        elif engine == "lmstudio":
            self.ai_key_label.configure(text="Model 名稱")
            self.ai_key_var.set(self.cfg.get("lmstudio_model", "google/gemma-4-e4b"))
            self.ai_key_entry.configure(show="")
        else:
            self.ai_key_label.configure(text="Gemini Key")
            self.ai_key_var.set(self.cfg.get("gemini_key", ""))
            self.ai_key_entry.configure(show="*")

    def _update_verify_btn(self, *_):
        """依輸入框是否有值切換「驗證」鈕外觀：有值→藍色醒目，空白→暗色。"""
        if not hasattr(self, "verify_btn"):
            return
        if self.ai_key_var.get().strip():
            self.verify_btn.configure(fg_color=ACCENT, hover_color=BLUE, text_color="white")
        else:
            self.verify_btn.configure(fg_color=SURFACE, hover_color=BORDER, text_color=SUBTEXT)

    def _verify_api(self):
        ai_engine = self.ai_engine_var.get()
        api_key = self.ai_key_var.get().strip()
        is_local = ai_engine in ("ollama", "lmstudio")
        if not api_key and not is_local:
            self.log_write(f"請先輸入 {_ENGINE_DISPLAY.get(ai_engine, ai_engine)} API Key")
            return
        self.verify_btn.configure(state="disabled", text="驗證中")
        self._set_status(f"正在驗證 {_ENGINE_DISPLAY.get(ai_engine, ai_engine)}...", ACCENT)

        def run():
            try:
                ok, msg = verify_engine(ai_engine, api_key)
            except Exception as e:
                ok, msg = False, str(e)
            self.root.after(0, lambda: self._verify_done(ai_engine, api_key, is_local, ok, msg))

        threading.Thread(target=run, daemon=True).start()

    def _verify_done(self, ai_engine, api_key, is_local, ok, msg):
        self.verify_btn.configure(state="normal", text="驗證")
        if ok:
            self.verify_btn.configure(fg_color=GREEN_DIM, text_color="white")
            self.log_write(f"✓ {msg}")
            self._set_status(f"✓ {msg}", GREEN)
            if not is_local:  # 驗證通過順手存下金鑰
                cfg_key = _ENGINE_CFG_KEY[ai_engine]
                save_config({cfg_key: api_key})
                self.cfg[cfg_key] = api_key
        else:
            self._update_verify_btn()
            self.log_write(f"✗ 驗證失敗：{msg}")
            self._set_status("驗證失敗", RED)

    def _on_cookies_change(self, choice):
        if choice == "Cookies: 選擇檔案...":
            self.pick_cookies_file()
        elif choice == "Cookies: 檔案":
            if not self.cookies_file_path or not os.path.exists(self.cookies_file_path):
                self.pick_cookies_file()

    def pick_cookies_file(self):
        path = filedialog.askopenfilename(
            title="選擇 Cookies 檔案",
            filetypes=[
                ("文字檔", "*.txt"),
                ("所有檔案", "*.*")
            ],
            parent=self.root
        )
        if path:
            self.cookies_file_path = path
            save_config({"cookies_file_path": path, "cookies_browser": "Cookies: 檔案"})
            self.cfg["cookies_file_path"] = path
            self.cfg["cookies_browser"] = "Cookies: 檔案"
            self.cookies_browser_var.set("Cookies: 檔案")
            self.log_write(f"已載入 Cookies 檔案：{os.path.basename(path)}")
        else:
            if not self.cookies_file_path or not os.path.exists(self.cookies_file_path):
                self.cookies_browser_var.set("Cookies: 無")

    def pick_audio(self):
        path = filedialog.askopenfilename(
            filetypes=[
                ("音訊/影片檔", "*.m4a *.mp3 *.wav *.mp4 *.mkv *.avi *.mov *.webm *.ogg *.flac"),
                ("音訊檔", "*.m4a *.mp3 *.wav *.ogg *.flac"),
                ("影片檔", "*.mp4 *.mkv *.avi *.mov *.webm"),
                ("所有檔案", "*.*"),
            ])
        if path:
            self.audio_var.set(path)
            self.audio_display_var.set(os.path.basename(path))
            self.audio_display_label.configure(text_color=ACCENT)
            ext = os.path.splitext(path)[1].lower()
            self.srt_var.set(ext in VIDEO_EXTS)
            if not self.out_var.get():
                self.out_var.set(os.path.dirname(path))
                self.out_display_var.set(os.path.dirname(path))

    def pick_output(self):
        path = filedialog.askdirectory()
        if path:
            self.out_var.set(path)
            self.out_display_var.set(path)

    def log_write(self, msg):
        self.log.configure(state="normal")
        self.log.insert(tk.END, msg + "\n")
        self.log.see(tk.END)
        self.log.configure(state="disabled")

    def update_timer(self):
        if self.timer_start is not None:
            elapsed = int(time.time() - self.timer_start)
            self.timer_var.set(f"{elapsed//3600:02d}:{(elapsed%3600)//60:02d}:{elapsed%60:02d}")
            self.timer_id = self.root.after(1000, self.update_timer)

    def _elapsed_str(self):
        """目前已經過的時間字串（HH:MM:SS）。"""
        if self.timer_start is None:
            return self.timer_var.get()
        el = int(time.time() - self.timer_start)
        return f"{el//3600:02d}:{(el%3600)//60:02d}:{el%60:02d}"

    def _stop_timer(self):
        """停止計時並把標籤凍結在最終時間。"""
        final = self._elapsed_str()
        if self.timer_id:
            self.root.after_cancel(self.timer_id)
            self.timer_id = None
        self.timer_start = None
        self.timer_var.set(final)
        return final

    # ── AI 進度動畫 ────────────────────────────────
    def _start_ai_progress(self):
        self.timer_start = time.time()
        self.update_timer()
        self.pct_var.set("")
        self._ai_anim_running = True
        self._ai_anim_val = 0.0
        self._ai_anim_dir = 1
        self._animate_ai_progress()

    def _animate_ai_progress(self):
        if not self._ai_anim_running:
            return
        self._ai_anim_val += self._ai_anim_dir * 0.025
        if self._ai_anim_val >= 1.0:
            self._ai_anim_val = 1.0
            self._ai_anim_dir = -1
        elif self._ai_anim_val <= 0.0:
            self._ai_anim_val = 0.0
            self._ai_anim_dir = 1
        self.progress.set(self._ai_anim_val)
        self.root.after(40, self._animate_ai_progress)

    def _stop_ai_progress(self, success=True):
        self._ai_anim_running = False
        self.progress.set(1.0 if success else 0.0)
        self._stop_timer()

    def _set_status(self, msg, color=SUBTEXT):
        self.status_label.configure(text=msg, text_color=color)

    def _short_name(self, name, maxlen=22):
        return name if len(name) <= maxlen else name[:maxlen - 1] + "…"

    # ── Poll & Transcription ───────────────────────
    def poll(self):
        try:
            while True:
                pct = self.prog_q.get_nowait()
                self.progress.set(pct / 100)
                self.pct_var.set(f"{pct}%")
        except Exception:
            pass
        try:
            while True:
                self.log_write(self.log_q.get_nowait())
        except Exception:
            pass
        try:
            status, data = self.result_q.get_nowait()
            # 先停錶並記下總耗時，避免後續任何步驟丟例外導致計時器停不下來
            elapsed = self._stop_timer()
            if status == "done":
                self.progress.set(1.0)
                self.pct_var.set("100%")
                self.log_write(f"完成！耗時 {elapsed}，儲存至：{data}")
                self._set_status(f"轉錄完成（耗時 {elapsed}）！已切到「逐字稿處理」分頁", GREEN)
                speakers = self._find_speakers(data)
                if speakers:
                    samples = self._speaker_samples(data, speakers)
                    dlg = SpeakerNameDialog(self.root, speakers, samples)
                    if dlg.result:
                        self._rename_speakers(data, dlg.result)
                self.transcript_path = data
                self._set_loaded_file(data)
                # 把第一步驟的「會議背景」自動帶進第二步驟的補充框（空才帶，不蓋掉手動輸入）
                try:
                    if not self.ai_context_box.get("1.0", "end").strip():
                        bg = self.bg_box.get("1.0", "end").strip()
                        if bg:
                            self.ai_context_box.insert("1.0", bg)
                except Exception:
                    pass
                _open_file(data)
                self.correct_btn.configure(state="normal", fg_color=BLUE, text_color="white")
                self.rename_spk_btn.configure(state="normal", fg_color=BLUE, text_color="white")
                self.notes_btn.configure(state="normal", fg_color=BLUE, text_color="white")
                self.export_srt_btn.configure(state="normal", fg_color=BLUE, text_color="white")
                # 轉錄完成自動帶使用者到下一步（AI 後製分頁）
                try:
                    self.tabs.set(self._TAB_AI)
                except Exception:
                    pass
            else:
                self.log_write(f"錯誤：{data}")
                self._set_status(f"發生錯誤（耗時 {elapsed}）", RED)
            self.finish()
            return
        except Exception:
            pass
        if self.process and self.process.is_alive():
            self.root.after(300, self.poll)
        else:
            self.finish()

    def start(self):
        audio = self.audio_var.get().strip()
        out_dir = self.out_var.get().strip()
        if not audio or not os.path.exists(audio):
            self._set_status("請先選擇音檔", RED)
            return
        if not out_dir:
            self._set_status("請選擇輸出資料夾", RED)
            return

        engine = self._engine_val()
        lang = LANG_MAP.get(self.lang_var.get())
        terms = self.terms_var.get().strip()
        prompt = _build_initial_prompt(terms)  # 地端 Whisper 用，已裁在 token 上限內

        write_srt = self.srt_var.get()
        if write_srt and engine == "whisper":
            switch = messagebox.askyesno(
                "建議切換引擎",
                "輸出字幕時建議使用 WhisperX 引擎，\n可提供逐字對齊、字幕時間戳更精準。\n\n要自動切換到 WhisperX 嗎？\n（選「否」繼續使用 Whisper）",
                parent=self.root,
            )
            if switch:
                self.engine_var.set(_STT_VAL2LABEL["whisperx"])
                self._on_engine_change()
                engine = "whisperx"

        if engine in ("whisperx", "whispercpp_diar"):
            hf_token = self.token_var.get().strip()
            if not hf_token:
                self._set_status("請輸入 HuggingFace Token", RED)
                return
            save_config({"hf_token": hf_token})
            num_speakers = self.speakers_var.get()

        # whisper.cpp 需要外部的 whisper-cli；沒裝就跳教學引導（含可一鍵複製的安裝指令），不硬跑
        if engine in ("whispercpp", "whispercpp_diar") and not _resolve_whispercpp_binary():
            self._show_whispercpp_guide()
            self._set_status("尚未安裝 whisper.cpp，請依視窗指示安裝", RED)
            return

        # 開始前就決定輸出路徑：同名先問「覆蓋還是另存」，把結果傳給 worker（三引擎一致）
        base = os.path.splitext(os.path.basename(audio))[0]
        out_path = self._resolve_output_path(os.path.join(out_dir, f"{base}_transcript.txt"))

        self.correct_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
        self.rename_spk_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
        self.notes_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
        self.transcript_path = None
        self.result_q = multiprocessing.Queue()
        self.log_q = multiprocessing.Queue()
        self.prog_q = multiprocessing.Queue()

        if engine == "whisperx":
            self.process = multiprocessing.Process(
                target=whisperx_worker,
                args=(audio, out_dir, self.model_var.get(), lang,
                      hf_token, num_speakers, prompt, write_srt,
                      self.result_q, self.log_q, self.prog_q, out_path),
                daemon=True)
        elif engine == "whispercpp":
            self.process = multiprocessing.Process(
                target=whispercpp_worker,
                args=(audio, out_dir, self.model_var.get(), lang, prompt, write_srt,
                      self.result_q, self.log_q, self.prog_q, out_path),
                daemon=True)
        elif engine == "whispercpp_diar":
            self.process = multiprocessing.Process(
                target=whispercpp_diar_worker,
                args=(audio, out_dir, self.model_var.get(), lang,
                      hf_token, num_speakers, prompt, write_srt,
                      self.result_q, self.log_q, self.prog_q, out_path),
                daemon=True)
        else:
            self.process = multiprocessing.Process(
                target=whisper_worker,
                args=(audio, out_dir, self.model_var.get(), lang, prompt, write_srt,
                      self.result_q, self.log_q, self.prog_q, out_path),
                daemon=True)

        self.process.start()
        self.progress.set(0)
        self.pct_var.set("0%")
        self.start_btn.configure(state="disabled", fg_color=GREEN_DIM)
        self.stop_btn.configure(state="normal")
        self._set_status("轉錄中，請稍候...", ACCENT)
        self.timer_start = time.time()
        self.update_timer()
        self.poll()

    def stop(self):
        if self.process and self.process.is_alive():
            self.process.terminate()
            self.process.join()
            self.log_write("已手動停止")
            self._set_status("已停止", SUBTEXT)
            self.finish()
        else:
            self.ai_cancelled = True
            self.log_write("正在取消，請稍候...")
            self._set_status("取消中...", SUBTEXT)

    def finish(self):
        self.start_btn.configure(state="normal", fg_color=GREEN)
        self.stop_btn.configure(state="disabled")
        self._stop_timer()
        self.process = None

    # ── Transcript helpers ─────────────────────────
    def load_transcript(self):
        path = filedialog.askopenfilename(
            filetypes=[("逐字稿 / Markdown 文件", "*.txt *.md *.markdown"),
                       ("Markdown 文件", "*.md *.markdown"),
                       ("所有檔案", "*.*")])
        if not path:
            return
        speakers = self._find_speakers(path)
        if speakers:
            samples = self._speaker_samples(path, speakers)
            dlg = SpeakerNameDialog(self.root, speakers, samples)
            if dlg.result:
                self._rename_speakers(path, dlg.result)
        self.transcript_path = path
        fname = os.path.basename(path)
        self._set_loaded_file(path)
        self.correct_btn.configure(state="normal", fg_color=BLUE, text_color="white")
        self.rename_spk_btn.configure(state="normal", fg_color=BLUE, text_color="white")
        self.notes_btn.configure(state="normal", fg_color=BLUE, text_color="white")
        self.export_srt_btn.configure(state="normal", fg_color=BLUE, text_color="white")
        self.log_write(f"已載入逐字稿：{fname}")
        self._set_status("已載入，請點選「校正逐字稿」或「產生摘要」", ACCENT)

    def _resolve_output_path(self, path):
        if not os.path.exists(path):
            return path
        overwrite = messagebox.askyesno(
            "檔案已存在",
            f"{os.path.basename(path)} 已存在，是否覆蓋？\n選「否」將自動另外命名。",
            parent=self.root,
        )
        if overwrite:
            return path
        base, ext = os.path.splitext(path)
        n = 2
        while os.path.exists(f"{base}_{n}{ext}"):
            n += 1
        return f"{base}_{n}{ext}"

    def _find_speakers(self, path):
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return sorted(set(re.findall(r"SPEAKER_\d+", content)))

    def _speaker_samples(self, path, speakers, max_len=90):
        """為每位說話者抓一句「最長的代表發言」，當作標名時的判斷提示。"""
        best = {spk: "" for spk in speakers}
        current = None
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.rstrip("\n")
                m = re.match(r"^(SPEAKER_\d+)：", line)
                if m:
                    current = m.group(1)
                    continue
                if current in best:
                    text = re.sub(r"^\[\d{2}:\d{2} --> \d{2}:\d{2}\]\s*", "", line).strip()
                    if len(text) > len(best[current]):
                        best[current] = text
        samples = {}
        for spk in speakers:
            t = best[spk]
            samples[spk] = (t[:max_len] + "…") if len(t) > max_len else t
        return samples

    def _rename_speakers(self, path, name_map):
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        for original, new_name in name_map.items():
            if original != new_name:
                content = content.replace(original, new_name)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

    def _find_party_labels(self, path):
        """找 AI 標註的「〔方別〕」標籤（如 主管／部屬），排除〔不確定〕。"""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        labels = []
        for m in re.findall(r"〔([^〕]+)〕", content):
            if m not in labels and m != "不確定":
                labels.append(m)
        return labels

    def _party_samples(self, path, labels, max_len=90):
        """為每個方別抓一句最長的代表發言當提示。"""
        best = {lb: "" for lb in labels}
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                m = re.match(r"^〔([^〕]+)〕", line.strip())
                if not m or m.group(1) not in best:
                    continue
                text = re.sub(r"^〔[^〕]+〕\s*\[\d{1,2}:\d{2} --> \d{1,2}:\d{2}\]\s*", "", line.strip()).strip()
                if len(text) > len(best[m.group(1)]):
                    best[m.group(1)] = text
        return {lb: (t[:max_len] + "…") if len(t) > max_len else t for lb, t in best.items()}

    def _rename_party_labels(self, path, name_map):
        """把〔舊名〕整批換成〔新名〕（只動方括號內，不碰內文裡的同字）。"""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        for old, new in name_map.items():
            if old != new:
                content = content.replace(f"〔{old}〕", f"〔{new}〕")
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

    # 「HH:MM:SS 講者」表頭格式（Plaud、以及已移除的舊雲端辨識都是這種）
    _TIMED_HDR = re.compile(r"^(\d{1,2}:\d{2}:\d{2})\s+(.+?)\s*$")

    def _find_timed_speakers(self, path):
        """找『時間戳 講者』表頭裡的講者標籤（如 Speaker 1、Speaker 2）。"""
        labels = []
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                m = self._TIMED_HDR.match(line.rstrip("\n"))
                if m:
                    lb = m.group(2).strip()
                    if lb and len(lb) <= 20 and lb not in labels:
                        labels.append(lb)
        return labels

    def _timed_speaker_samples(self, path, labels, max_len=90):
        best = {lb: "" for lb in labels}
        current = None
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.rstrip("\n")
                m = self._TIMED_HDR.match(line)
                if m and m.group(2).strip() in best:
                    current = m.group(2).strip()
                    continue
                if current in best:
                    t = line.strip()
                    if len(t) > len(best[current]):
                        best[current] = t
        return {lb: (t[:max_len] + "…") if len(t) > max_len else t for lb, t in best.items()}

    def _rename_timed_speakers(self, path, name_map):
        """只改表頭行的講者名（行內文字不動）。"""
        out = []
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                m = self._TIMED_HDR.match(line.rstrip("\n"))
                if m and m.group(2).strip() in name_map and name_map[m.group(2).strip()] != m.group(2).strip():
                    out.append(f"{m.group(1)} {name_map[m.group(2).strip()]}\n")
                else:
                    out.append(line)
        with open(path, "w", encoding="utf-8") as f:
            f.writelines(out)

    def rename_speakers(self):
        """事後修正說話者/說話方名稱：支援 WhisperX 的 SPEAKER_xx、AI 標的〔方別〕、
        以及 Plaud／舊雲端的『HH:MM:SS 講者』表頭格式。"""
        if not self.transcript_path:
            return
        path = self.transcript_path
        spk = self._find_speakers(path)
        if spk:
            dlg = SpeakerNameDialog(self.root, spk, self._speaker_samples(path, spk))
            if dlg.result:
                self._rename_speakers(path, dlg.result)
                self.log_write("已更新說話者名稱")
                self._set_status("已更新說話者名稱！", GREEN)
                _open_file(path)
            return
        labels = self._find_party_labels(path)
        if labels:
            dlg = SpeakerNameDialog(self.root, labels, self._party_samples(path, labels))
            if dlg.result:
                self._rename_party_labels(path, dlg.result)
                self.log_write("已更新說話方名稱")
                self._set_status("已更新說話方名稱！", GREEN)
                _open_file(path)
            return
        timed = self._find_timed_speakers(path)
        if timed:
            dlg = SpeakerNameDialog(self.root, timed, self._timed_speaker_samples(path, timed))
            if dlg.result:
                self._rename_timed_speakers(path, dlg.result)
                self.log_write("已更新說話者名稱")
                self._set_status("已更新說話者名稱！", GREEN)
                _open_file(path)
            return
        self.log_write("這份逐字稿沒有可重新命名的說話人／說話方標籤")
        self._set_status("沒有可重新命名的標籤", SUBTEXT)

    # ── AI operations ──────────────────────────────
    def _ai_context(self):
        """AI 校正/摘要用的情境：合併第一步驟的關鍵術語＋會議背景，與第二步驟的補充說明。
        會議背景在轉錄完成後會自動帶入第二步驟的補充框，故以第二步驟的內容為主、背景為備。"""
        parts = []
        terms = self.terms_var.get().strip()
        if terms:
            parts.append(f"關鍵術語/人名：{terms}")
        try:
            bg1 = self.bg_box.get("1.0", "end").strip()
        except Exception:
            bg1 = ""
        try:
            t2 = self.ai_context_box.get("1.0", "end").strip()
        except Exception:
            t2 = ""
        if t2:
            parts.append(t2)
        elif bg1:
            parts.append(bg1)
        return "　".join(parts)

    def correct_transcript(self):
        if not self.transcript_path:
            return
        ai_engine = self.ai_engine_var.get()
        api_key = self.ai_key_var.get().strip()
        if not api_key and ai_engine not in ("ollama", "lmstudio"):
            self.log_write(f"請輸入 {_ENGINE_DISPLAY.get(ai_engine, ai_engine)} API Key")
            return
        # 稿子已有任何說話人標記（WhisperX SPEAKER_xx／AI〔方別〕／Plaud 的 HH:MM:SS 講者）
        # → 只校正；完全沒有（whisper/whisper.cpp 原稿）→ 同一次 API 順手標方
        p = self.transcript_path
        do_label = not (self._find_speakers(p) or self._find_party_labels(p) or self._find_timed_speakers(p))
        context = self._ai_context()
        if do_label:
            # 要標方時先問「誰是誰」：預填現有情境，可補人名；取消就不跑
            dlg = PartyContextDialog(self.root, prefill=context)
            if dlg.result is None:
                return
            context = dlg.result
        self._correct_did_label = do_label

        base = self.transcript_path.rsplit("_transcript", 1)[0]
        out_path = self._resolve_output_path(base + "_transcript_corrected.txt")
        cfg_key = _ENGINE_CFG_KEY[ai_engine]
        save_config({cfg_key: api_key})
        self.cfg[cfg_key] = api_key

        self.ai_cancelled = False
        self.correct_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
        self.notes_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
        self.stop_btn.configure(state="normal")
        action = "校正並標註說話方" if do_label else "校正逐字稿"
        self.log_write(f"正在{action}（{_ENGINE_DISPLAY.get(ai_engine, ai_engine)}）...")
        self._set_status(f"{action}中...", ACCENT)
        self._start_ai_progress()

        def run():
            try:
                result = correct_transcript(self.transcript_path, ai_engine, api_key, out_path,
                                            context=context, label_parties=do_label)
                self.root.after(0, lambda: self._correct_done(result))
            except Exception as e:
                err = str(e)
                self.root.after(0, lambda: self._correct_error(err))

        threading.Thread(target=run, daemon=True).start()

    def _correct_done(self, out_path):
        self._stop_ai_progress(success=not self.ai_cancelled)
        self._reset_ai_buttons()
        if self.ai_cancelled:
            self.log_write("已取消")
            self._set_status("已取消", SUBTEXT)
            return
        self.transcript_path = out_path
        self._set_loaded_file(out_path)
        action = "校正＋標註說話方" if getattr(self, "_correct_did_label", False) else "校正"
        self.log_write(f"{action}完成：{os.path.basename(out_path)}")
        self._set_status(f"{action}完成！已自動開啟，可繼續點選「產生摘要」", GREEN)
        _open_file(out_path)

    def _correct_error(self, err):
        self._stop_ai_progress(success=False)
        self._reset_ai_buttons()
        self.log_write(f"校正錯誤：{err}")
        self._set_status("校正失敗", RED)

    def _set_loaded_file(self, path):
        """更新「目前處理檔案」醒目標示。"""
        self.loaded_file_var.set(f"目前處理：{os.path.basename(path)}")
        try:
            self.loaded_file_label.configure(text_color=ACCENT)
        except Exception:
            pass

    def _show_whispercpp_guide(self):
        """選了 whisper.cpp 但沒裝 whisper-cli 時，跳教學引導（含可一鍵複製的安裝指令）。"""
        is_mac = sys.platform == "darwin"
        cmd = "brew install whisper-cpp"
        win = ctk.CTkToplevel(self.root)
        win.title("要用 whisper.cpp，需要先安裝一次")
        win.configure(fg_color=BG)
        win.resizable(False, False)
        win.grab_set()
        ctk.CTkLabel(
            win, justify="left", fg_color="transparent", text_color=TEXT,
            font=ctk.CTkFont(FONT_UI, 15, weight="bold"),
            text="「whisper.cpp」是給 MacBook Air 等較輕薄機器用的辨識引擎",
        ).pack(anchor="w", padx=22, pady=(20, 2))
        ctk.CTkLabel(
            win, justify="left", fg_color="transparent", text_color=SUBTEXT,
            font=ctk.CTkFont(FONT_UI, 13), wraplength=460,
            text="第一次使用前需要安裝一次（之後就不用再裝）。轉錄用的模型會在第一次轉錄時自動下載，不用手動處理。",
        ).pack(anchor="w", padx=22, pady=(0, 12))

        if is_mac:
            ctk.CTkLabel(
                win, justify="left", fg_color="transparent", text_color=TEXT,
                font=ctk.CTkFont(FONT_UI, 13),
                text="① 打開「終端機」　② 貼上並執行下面這行：",
            ).pack(anchor="w", padx=22, pady=(0, 6))
            box = ctk.CTkFrame(win, fg_color=SURFACE, corner_radius=8)
            box.pack(fill="x", padx=22)
            ctk.CTkLabel(box, text=cmd, fg_color="transparent", text_color=ACCENT,
                         font=ctk.CTkFont("Menlo", 14, weight="bold"),
                         anchor="w").pack(side="left", padx=12, pady=10)

            def _copy():
                self.root.clipboard_clear()
                self.root.clipboard_append(cmd)
                copy_btn.configure(text="已複製 ✓")

            copy_btn = ctk.CTkButton(box, text="複製指令", width=92, command=_copy,
                                     fg_color=ACCENT, hover_color="#2980B9", text_color="white",
                                     font=ctk.CTkFont(FONT_UI, 13))
            copy_btn.pack(side="right", padx=8, pady=7)
            ctk.CTkLabel(
                win, justify="left", fg_color="transparent", text_color=SUBTEXT,
                font=ctk.CTkFont(FONT_UI, 12), wraplength=460,
                text="③ 裝完回到 VoxLog，再按一次「開始轉錄」即可。",
            ).pack(anchor="w", padx=22, pady=(10, 4))
        else:
            ctk.CTkLabel(
                win, justify="left", fg_color="transparent", text_color=SUBTEXT,
                font=ctk.CTkFont(FONT_UI, 13), wraplength=460,
                text="Windows 版的 whisper.cpp 執行檔原本已內建在專案的 bin\\whispercpp-win\\ 裡，"
                     "正常不會看到這個訊息。若出現，多半是檔案不齊，請重新執行一次 git pull（或重新下載專案）。",
            ).pack(anchor="w", padx=22, pady=(0, 8))

        ctk.CTkButton(win, text="知道了", width=110, command=win.destroy,
                      fg_color=GREEN, hover_color="#219A52", text_color="white",
                      font=ctk.CTkFont(FONT_UI, 13)).pack(pady=(10, 20))
        win.transient(self.root)

    def generate_notes(self):
        if not self.transcript_path:
            return
        ai_engine = self.ai_engine_var.get()
        api_key = self.ai_key_var.get().strip()
        if not api_key and ai_engine not in ("ollama", "lmstudio"):
            self.log_write(f"請輸入 {_ENGINE_DISPLAY.get(ai_engine, ai_engine)} API Key")
            return
        base = _clean_basename(self.transcript_path)
        default_out = os.path.join(os.path.dirname(self.transcript_path), f"{base}_摘要.docx")
        out_path = self._resolve_output_path(default_out)
        cfg_key = _ENGINE_CFG_KEY[ai_engine]
        save_config({cfg_key: api_key})
        self.cfg[cfg_key] = api_key

        self.ai_cancelled = False
        self.correct_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
        self.notes_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
        self.stop_btn.configure(state="normal")
        self.log_write(f"正在產生摘要（{_ENGINE_DISPLAY.get(ai_engine, ai_engine)}）...")
        self._set_status("產生摘要中...", ACCENT)
        self._start_ai_progress()

        def run():
            try:
                result = generate_summary(self.transcript_path, ai_engine, api_key, out_path, context=self._ai_context())
                self.root.after(0, lambda: self._notes_done(result))
            except Exception as e:
                err = str(e)
                self.root.after(0, lambda: self._notes_error(err))

        self.notes_thread = threading.Thread(target=run, daemon=True)
        self.notes_thread.start()

    def export_srt(self):
        if not self.transcript_path:
            return
        base = self.transcript_path.rsplit(".", 1)[0]
        for suffix in ("_transcript_corrected", "_transcript"):
            if base.endswith(suffix):
                base = base[:-len(suffix)]
                break
        out_path = self._resolve_output_path(base + ".srt")

        ai_engine = self.ai_engine_var.get()
        api_key = self.ai_key_var.get().strip()

        if api_key:
            self.ai_cancelled = False
            self.export_srt_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
            self.correct_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
            self.notes_btn.configure(state="disabled", fg_color=BLUE_DIM, text_color=SUBTEXT)
            self.stop_btn.configure(state="normal")
            self.log_write(f"正在用 AI 切割字幕（{_ENGINE_DISPLAY.get(ai_engine, ai_engine)}）...")
            self._set_status("AI 切割字幕中...", ACCENT)
            self._start_ai_progress()

            srt_source = base + ".srt"
            use_srt = os.path.exists(srt_source)
            if use_srt:
                self.log_write("找到原始 SRT，使用毫秒精準時間戳...")
            else:
                self.log_write("未找到原始 SRT，改用 TXT 時間戳（整數秒）...")

            def run():
                try:
                    if use_srt:
                        segments = _parse_srt(srt_source)
                        entries = _ai_cut_segments(segments, ai_engine, api_key)
                        _write_srt(entries, out_path)
                        result = out_path
                    else:
                        result = txt_to_srt_ai(self.transcript_path, ai_engine, api_key, out_path)
                    self.root.after(0, lambda: self._export_srt_done(result))
                except Exception as e:
                    err = str(e)
                    self.root.after(0, lambda: self._export_srt_error(err))

            threading.Thread(target=run, daemon=True).start()
        else:
            try:
                txt_to_srt(self.transcript_path, out_path)
                self.log_write(f"SRT 已匯出：{os.path.basename(out_path)}")
                self._set_status("SRT 匯出完成！", GREEN)
            except Exception as e:
                self.log_write(f"SRT 匯出錯誤：{e}")
                self._set_status("SRT 匯出失敗", RED)

    def _export_srt_done(self, out_path):
        self._stop_ai_progress(success=not self.ai_cancelled)
        self._reset_ai_buttons()
        if self.ai_cancelled:
            self.log_write("SRT 匯出已取消")
            self._set_status("已取消", SUBTEXT)
            return
        self.log_write(f"SRT 已匯出：{os.path.basename(out_path)}")
        self._set_status("SRT 匯出完成！", GREEN)

    def _export_srt_error(self, err):
        self._stop_ai_progress(success=False)
        self._reset_ai_buttons()
        self.log_write(f"SRT 匯出錯誤：{err}")
        self._set_status("SRT 匯出失敗", RED)

    def _notes_done(self, out_path):
        self._stop_ai_progress(success=not self.ai_cancelled)
        self._reset_ai_buttons()
        if self.ai_cancelled:
            self.log_write("產生摘要已取消")
            self._set_status("已取消", SUBTEXT)
            return
        self.log_write(f"摘要已儲存：{os.path.basename(out_path)}")
        self._set_status("摘要完成！已自動開啟", GREEN)
        _open_file(out_path)

    def _notes_error(self, err):
        self._stop_ai_progress(success=False)
        self._reset_ai_buttons()
        self.log_write(f"摘要錯誤：{err}")
        self._set_status("摘要失敗", RED)

    def _reset_ai_buttons(self):
        self.stop_btn.configure(state="disabled")
        self.correct_btn.configure(state="normal", fg_color=BLUE, text_color="white")
        self.rename_spk_btn.configure(state="normal", fg_color=BLUE, text_color="white")
        self.notes_btn.configure(state="normal", fg_color=BLUE, text_color="white")
        self.export_srt_btn.configure(state="normal", fg_color=BLUE, text_color="white")


def _is_masked_widget(w):
    """遮罩欄位（show="*"）= API key / token，使用者一定整段貼上。"""
    try:
        return bool(w.cget("show"))
    except Exception:
        return False


def _widget_paste(w):
    """手動貼上：遮罩欄位整段取代＋去空白，其餘插入游標處。回傳 'break'。"""
    try:
        text = w.clipboard_get()
    except Exception:
        return "break"
    if _is_masked_widget(w):
        try:
            w.delete(0, "end")
            w.insert(0, text.strip())
            return "break"
        except Exception:
            pass
    try:
        w.delete("sel.first", "sel.last")   # 有選取就先覆蓋
    except Exception:
        pass
    try:
        w.insert("insert", text)
    except Exception:
        try:
            w.event_generate("<<Paste>>")
        except Exception:
            pass
    return "break"


def _widget_select_all(w):
    try:
        w.select_range(0, "end")            # Entry / CTkEntry
        w.icursor("end")
    except Exception:
        try:
            w.tag_add("sel", "1.0", "end")  # Text
        except Exception:
            pass
    return "break"


def _enable_mac_clipboard(root):
    """macOS：tkinter 的 Cmd+C/V/X/A 在不同焦點時機常失靈（平台老問題）。
    這裡用兩道保險：① 攔截 Cmd 鍵手動處理；② 右鍵選單作為 100% 可靠的後備。"""
    if sys.platform != "darwin":
        return

    def _fire(event_name):
        def handler(event):
            try:
                event.widget.event_generate(event_name)
            except Exception:
                pass
            return "break"
        return handler

    root.bind_all("<Command-c>", _fire("<<Copy>>"))
    root.bind_all("<Command-v>", lambda e: _widget_paste(e.widget))
    root.bind_all("<Command-x>", _fire("<<Cut>>"))
    root.bind_all("<Command-a>", lambda e: _widget_select_all(e.widget))

    # 右鍵選單：不依賴快捷鍵，永遠可用
    def _popup(event):
        w = event.widget
        try:
            w.focus_set()
        except Exception:
            pass
        menu = tk.Menu(w, tearoff=0)
        menu.add_command(label="剪下", command=lambda: w.event_generate("<<Cut>>"))
        menu.add_command(label="複製", command=lambda: w.event_generate("<<Copy>>"))
        menu.add_command(label="貼上", command=lambda: _widget_paste(w))
        menu.add_separator()
        menu.add_command(label="全選", command=lambda: _widget_select_all(w))
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()
        return "break"

    # 綁在 Entry / Text class 上，連之後新建的欄位、對話框都涵蓋
    for cls in ("Entry", "Text"):
        root.bind_class(cls, "<Button-2>", _popup, add="+")          # 部分滑鼠的右鍵
        root.bind_class(cls, "<Button-3>", _popup, add="+")          # 多數滑鼠的右鍵
        root.bind_class(cls, "<Control-Button-1>", _popup, add="+")  # Ctrl+左鍵 = 右鍵


if __name__ == "__main__":
    multiprocessing.freeze_support()
    root = ctk.CTk()
    app = TranscribeApp(root)
    _enable_mac_clipboard(root)
    root.mainloop()
